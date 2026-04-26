import logging
import csv
import io
import json
import re
import unicodedata
from dataclasses import dataclass, field
from datetime import datetime, timezone
from decimal import Decimal
from pathlib import Path
from time import perf_counter
from typing import Any, Callable
from uuid import UUID, uuid4

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.core.constants import ExecutionStatus
from app.core.exceptions import AppException
from app.core.log_context import bind_log_context, reset_log_context
from app.integrations.providers.base import ProviderExecutionResult
from app.integrations.queue.dispatcher import enqueue_execution_job
from app.models.operational import (
    DjangoAiApiToken,
    DjangoAiApiTokenPermission,
    DjangoAiAutomationExecutionSetting,
    DjangoAiAuditLog,
    DjangoAiExecutionInputFile,
    DjangoAiQueueJob,
)
from app.repositories.operational import (
    AuditLogRepository,
    AutomationExecutionSettingsRepository,
    ExecutionInputFileRepository,
    QueueJobRepository,
    RequestFileRepository,
)
from app.repositories.shared import SharedAnalysisRepository, SharedExecutionRepository
from app.services.execution_engine import (
    ALLOWED_INPUT_ROLES,
    INPUT_ROLE_CONTEXT,
    INPUT_ROLE_PRIMARY,
    EngineExecutionInput,
    EngineExecutionPlan,
    ExecutionOutputContract,
    ExecutionFormatterStrategy,
    ExecutionInputType,
    ExecutionOutputPolicy,
    ExecutionOutputType,
    ExecutionParserStrategy,
    ExecutionResponseParser,
    ExecutionStrategyEngine,
)
from app.services.execution_output_contract import ExecutionOutputContractResolver
from app.services.execution_output_pipeline import ExecutionResultFormatter, ExecutionResultNormalizer
from app.services.execution_tabular_prompt_strategy import TabularPromptStrategyResolver
from app.services.execution_observability import (
    ExecutionErrorDiagnostic,
    classify_execution_error,
    summarize_processing_inputs,
    summarize_processing_plan,
)
from app.services.file_service import FileService
from app.services.providers.http_client_utils import (
    sanitize_provider_debug_payload,
    summarize_provider_error_message,
)
from app.services.provider_service import ProviderRuntimeSelection, ProviderService
from app.services.shared.automation_runtime_resolver import AutomationRuntimeResolverService
from app.services.token_service import check_token_permission
from app.services.usage_service import UsageService

settings = get_settings()
logger = logging.getLogger(__name__)

RETRYABLE_PROVIDER_STATUS_CODES = {408, 409, 425, 429, 500, 502, 503, 504}
RETRYABLE_ERROR_CODES = {"provider_timeout", "provider_network_error"}
TABULAR_EXTENSIONS = {".xlsx", ".xls", ".csv"}
TEXTUAL_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".txt",
    ".md",
    ".json",
    ".xml",
    ".html",
    ".htm",
    ".rtf",
    ".log",
    ".yaml",
    ".yml",
}
TABULAR_MIME_HINTS = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
    "text/csv",
    "application/csv",
}
TEXTUAL_MIME_HINTS = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "application/json",
    "application/xml",
}
LEGACY_XLS_EXTENSION = ".xls"
HARD_LIMIT_ERROR_CODES = {
    "execution_rows_hard_limit_exceeded",
    "provider_calls_hard_limit_exceeded",
    "text_chunks_hard_limit_exceeded",
    "tabular_row_size_hard_limit_exceeded",
    "execution_time_hard_limit_exceeded",
    "job_retries_hard_limit_exceeded",
}
PROFILE_LIMIT_ERROR_CODES = {
    "execution_rows_profile_limit_exceeded",
    "provider_calls_profile_limit_exceeded",
    "text_chunks_profile_limit_exceeded",
    "tabular_row_size_profile_limit_exceeded",
    "execution_time_profile_limit_exceeded",
}
FATAL_TABULAR_ERROR_CODES = {
    "cost_limit_exceeded",
    "prompt_token_limit_exceeded",
    "prompt_placeholder_unresolved",
    *HARD_LIMIT_ERROR_CODES,
    *PROFILE_LIMIT_ERROR_CODES,
}

CONTEXT_STRUCTURED_EXTENSIONS = {".json", ".xml", ".yaml", ".yml", ".csv", ".tsv"}
CONTEXT_RAW_EXTENSIONS = {".txt", ".md", ".log", ".pdf", ".docx", ".doc", ".html", ".htm", ".rtf"}
PROMPT_SECTION_INSTRUCTION = "[INSTRUCAO]"
PROMPT_SECTION_ROW_DATA = "[DADOS DA LINHA]"
PROMPT_SECTION_CONTEXT = "[CONTEXTO]"

PROFILE_STANDARD = "standard"
PROFILE_HEAVY = "heavy"
PROFILE_EXTENDED = "extended"
KNOWN_EXECUTION_PROFILES = {PROFILE_STANDARD, PROFILE_HEAVY, PROFILE_EXTENDED}

LIMIT_KEY_MAX_EXECUTION_ROWS = "max_execution_rows"
LIMIT_KEY_MAX_PROVIDER_CALLS = "max_provider_calls"
LIMIT_KEY_MAX_TEXT_CHUNKS = "max_text_chunks"
LIMIT_KEY_MAX_TABULAR_ROW_CHARACTERS = "max_tabular_row_characters"
LIMIT_KEY_MAX_EXECUTION_SECONDS = "max_execution_seconds"
LIMIT_KEY_MAX_CONTEXT_CHARACTERS = "max_context_characters"
LIMIT_KEY_MAX_CONTEXT_FILE_CHARACTERS = "max_context_file_characters"
LIMIT_KEY_MAX_PROMPT_CHARACTERS = "max_prompt_characters"


@dataclass(slots=True)
class ExecutionCreateResult:
    execution_id: UUID
    queue_job_id: UUID
    status: ExecutionStatus


@dataclass(slots=True)
class ExecutionStatusResult:
    execution_id: UUID
    status: ExecutionStatus
    progress: int | None
    started_at: datetime | None
    finished_at: datetime | None
    error_message: str | None
    created_at: datetime


@dataclass(slots=True)
class ExecutionInputSelection:
    request_file_id: UUID
    role: str
    order_index: int
    file_name: str | None = None


@dataclass(slots=True)
class ExecutionInputResult:
    request_file_id: UUID
    file_name: str | None
    role: str
    order_index: int
    source: str


@dataclass(slots=True)
class ProcessedOutput:
    content: bytes
    file_name: str
    mime_type: str
    total_input_tokens: int
    total_output_tokens: int
    total_cost: Decimal
    providers_used: set[str]
    models_used: set[str]
    provider_calls: int
    processing_summary: dict[str, Any]
    auxiliary_files: list["GeneratedExecutionFile"] = field(default_factory=list)


@dataclass(slots=True, frozen=True)
class ExecutionProgressUpdate:
    phase: str
    status_message: str
    processed_rows: int | None = None
    total_rows: int | None = None
    current_row: int | None = None
    processed_chunks: int | None = None
    total_chunks: int | None = None


@dataclass(slots=True)
class GeneratedExecutionFile:
    file_type: str
    file_name: str
    mime_type: str
    content: bytes


@dataclass(slots=True, frozen=True)
class ExecutionOperationalProfile:
    name: str
    source: str
    source_details: dict[str, Any]
    persisted_overrides: dict[str, int]
    max_execution_rows: int
    max_provider_calls: int
    max_text_chunks: int
    max_tabular_row_characters: int
    max_execution_seconds: int
    max_context_characters: int
    max_context_file_characters: int
    max_prompt_characters: int
    hard_clamped_fields: tuple[str, ...]
    hard_clamp_details: dict[str, dict[str, int]]

    def to_limits_dict(self) -> dict[str, int]:
        return {
            LIMIT_KEY_MAX_EXECUTION_ROWS: int(self.max_execution_rows),
            LIMIT_KEY_MAX_PROVIDER_CALLS: int(self.max_provider_calls),
            LIMIT_KEY_MAX_TEXT_CHUNKS: int(self.max_text_chunks),
            LIMIT_KEY_MAX_TABULAR_ROW_CHARACTERS: int(self.max_tabular_row_characters),
            LIMIT_KEY_MAX_EXECUTION_SECONDS: int(self.max_execution_seconds),
            LIMIT_KEY_MAX_CONTEXT_CHARACTERS: int(self.max_context_characters),
            LIMIT_KEY_MAX_CONTEXT_FILE_CHARACTERS: int(self.max_context_file_characters),
            LIMIT_KEY_MAX_PROMPT_CHARACTERS: int(self.max_prompt_characters),
        }

    def is_hard_clamped(self, limit_key: str) -> bool:
        return limit_key in self.hard_clamped_fields


class ExecutionService:
    def __init__(
        self,
        *,
        operational_session: Session,
        shared_session: Session,
    ) -> None:
        self.operational_session = operational_session
        self.shared_session = shared_session
        self.request_files = RequestFileRepository(operational_session)
        self.execution_inputs = ExecutionInputFileRepository(operational_session)
        self.execution_profile_settings = AutomationExecutionSettingsRepository(operational_session)
        self.queue_jobs = QueueJobRepository(operational_session)
        self.audit_logs = AuditLogRepository(operational_session)
        self.shared_analysis = SharedAnalysisRepository(shared_session)
        self.shared_executions = SharedExecutionRepository(shared_session)
        self.runtime_resolver = AutomationRuntimeResolverService(
            shared_session=shared_session,
            operational_session=operational_session,
        )
        self.provider_service = ProviderService(operational_session)
        self.usage_service = UsageService(operational_session)
        self.file_service = FileService(
            operational_session=operational_session,
            shared_session=shared_session,
        )
        self.strategy_engine = ExecutionStrategyEngine(
            tabular_extensions=TABULAR_EXTENSIONS,
            textual_extensions=TEXTUAL_EXTENSIONS,
            tabular_mime_hints=TABULAR_MIME_HINTS,
            textual_mime_hints=TEXTUAL_MIME_HINTS,
        )
        self.output_contract_resolver = ExecutionOutputContractResolver()
        self.response_parser = ExecutionResponseParser()
        self.result_normalizer = ExecutionResultNormalizer()
        self.result_formatter = ExecutionResultFormatter()
        self.tabular_prompt_strategy_resolver = TabularPromptStrategyResolver()
        self.output_policy = ExecutionOutputPolicy()

    def _log_execution_phase(
        self,
        *,
        phase: str,
        message: str,
        level: str = "info",
        **extra_fields: Any,
    ) -> None:
        payload = {
            "event": "execution_phase",
            "phase": phase,
            **extra_fields,
        }
        log_method = getattr(logger, level, logger.info)
        log_method(message, extra=payload)

    @staticmethod
    def _safe_progress_int(value: Any) -> int | None:
        if value is None:
            return None
        try:
            return int(value)
        except (TypeError, ValueError):
            return None

    def _notify_progress_update(
        self,
        *,
        progress_callback: Callable[[ExecutionProgressUpdate], None] | None,
        phase: str,
        status_message: str,
        processed_rows: int | None = None,
        total_rows: int | None = None,
        current_row: int | None = None,
        processed_chunks: int | None = None,
        total_chunks: int | None = None,
    ) -> None:
        if progress_callback is None:
            return
        try:
            progress_callback(
                ExecutionProgressUpdate(
                    phase=str(phase or "").strip().lower() or "running_model",
                    status_message=str(status_message or "").strip() or "Processando execucao.",
                    processed_rows=self._safe_progress_int(processed_rows),
                    total_rows=self._safe_progress_int(total_rows),
                    current_row=self._safe_progress_int(current_row),
                    processed_chunks=self._safe_progress_int(processed_chunks),
                    total_chunks=self._safe_progress_int(total_chunks),
                )
            )
        except Exception:
            logger.warning("Execution progress callback failed.", exc_info=True)

    @staticmethod
    def _safe_hard_limit(value: int, *, fallback: int = 1) -> int:
        try:
            normalized = int(value)
        except (TypeError, ValueError):
            normalized = fallback
        return max(normalized, fallback)

    @staticmethod
    def _normalize_profile_name(value: str | None) -> str:
        return str(value or "").strip().lower()

    def _hard_limit_ceilings(self) -> dict[str, int]:
        return {
            LIMIT_KEY_MAX_EXECUTION_ROWS: self._safe_hard_limit(settings.max_execution_rows_hard_limit, fallback=1),
            LIMIT_KEY_MAX_PROVIDER_CALLS: self._safe_hard_limit(settings.max_provider_calls_hard_limit, fallback=1),
            LIMIT_KEY_MAX_TEXT_CHUNKS: self._safe_hard_limit(settings.max_text_chunks_hard_limit, fallback=1),
            LIMIT_KEY_MAX_TABULAR_ROW_CHARACTERS: self._safe_hard_limit(
                settings.max_tabular_row_characters_hard_limit,
                fallback=1,
            ),
            LIMIT_KEY_MAX_EXECUTION_SECONDS: self._safe_hard_limit(settings.max_execution_seconds_hard_limit, fallback=1),
            LIMIT_KEY_MAX_CONTEXT_CHARACTERS: self._safe_hard_limit(settings.max_context_characters, fallback=1),
            LIMIT_KEY_MAX_CONTEXT_FILE_CHARACTERS: self._safe_hard_limit(settings.max_context_file_characters, fallback=1),
            LIMIT_KEY_MAX_PROMPT_CHARACTERS: self._safe_hard_limit(settings.max_prompt_characters, fallback=1),
        }

    def _profile_limits_from_settings(self, *, profile_name: str) -> dict[str, int]:
        profile_prefix = f"execution_profile_{profile_name}_"
        setting_suffix_by_limit = {
            LIMIT_KEY_MAX_EXECUTION_ROWS: "max_execution_rows",
            LIMIT_KEY_MAX_PROVIDER_CALLS: "max_provider_calls",
            LIMIT_KEY_MAX_TEXT_CHUNKS: "max_text_chunks",
            LIMIT_KEY_MAX_TABULAR_ROW_CHARACTERS: "max_tabular_row_characters",
            LIMIT_KEY_MAX_EXECUTION_SECONDS: "max_execution_seconds",
            LIMIT_KEY_MAX_CONTEXT_CHARACTERS: "max_context_characters",
            LIMIT_KEY_MAX_CONTEXT_FILE_CHARACTERS: "max_context_file_characters",
            LIMIT_KEY_MAX_PROMPT_CHARACTERS: "max_prompt_characters",
        }
        hard_ceilings = self._hard_limit_ceilings()
        resolved: dict[str, int] = {}
        for limit_key, setting_suffix in setting_suffix_by_limit.items():
            setting_name = f"{profile_prefix}{setting_suffix}"
            raw_value = getattr(settings, setting_name, None)
            fallback_value = hard_ceilings[limit_key]
            resolved[limit_key] = self._safe_hard_limit(raw_value if raw_value is not None else fallback_value, fallback=1)
        return resolved

    @staticmethod
    def _persisted_limit_field_by_key() -> dict[str, str]:
        return {
            LIMIT_KEY_MAX_EXECUTION_ROWS: "max_execution_rows",
            LIMIT_KEY_MAX_PROVIDER_CALLS: "max_provider_calls",
            LIMIT_KEY_MAX_TEXT_CHUNKS: "max_text_chunks",
            LIMIT_KEY_MAX_TABULAR_ROW_CHARACTERS: "max_tabular_row_characters",
            LIMIT_KEY_MAX_EXECUTION_SECONDS: "max_execution_seconds",
            LIMIT_KEY_MAX_CONTEXT_CHARACTERS: "max_context_characters",
            LIMIT_KEY_MAX_CONTEXT_FILE_CHARACTERS: "max_context_file_characters",
            LIMIT_KEY_MAX_PROMPT_CHARACTERS: "max_prompt_characters",
        }

    def _extract_persisted_overrides(
        self,
        *,
        persisted_setting: DjangoAiAutomationExecutionSetting,
    ) -> dict[str, int]:
        overrides: dict[str, int] = {}
        for limit_key, field_name in self._persisted_limit_field_by_key().items():
            raw_value = getattr(persisted_setting, field_name, None)
            if raw_value is None:
                continue
            overrides[limit_key] = self._safe_hard_limit(raw_value, fallback=1)
        return overrides

    @staticmethod
    def _merge_profile_limits(
        *,
        base_limits: dict[str, int],
        override_limits: dict[str, int],
    ) -> dict[str, int]:
        merged = dict(base_limits)
        for limit_key, limit_value in override_limits.items():
            if limit_key in merged:
                merged[limit_key] = int(limit_value)
        return merged

    @staticmethod
    def _apply_hard_limit_ceilings(
        *,
        configured_limits: dict[str, int],
        hard_ceilings: dict[str, int],
    ) -> tuple[dict[str, int], tuple[str, ...], dict[str, dict[str, int]]]:
        effective_limits: dict[str, int] = {}
        hard_clamped_fields: list[str] = []
        hard_clamp_details: dict[str, dict[str, int]] = {}
        for limit_key, configured_value in configured_limits.items():
            ceiling = hard_ceilings[limit_key]
            effective_value = min(int(configured_value), int(ceiling))
            effective_limits[limit_key] = effective_value
            if configured_value > ceiling:
                hard_clamped_fields.append(limit_key)
                hard_clamp_details[limit_key] = {
                    "profile_value": int(configured_value),
                    "hard_limit": int(ceiling),
                }
        return effective_limits, tuple(sorted(hard_clamped_fields)), hard_clamp_details

    def _resolve_execution_profile(self, *, automation_id: UUID) -> ExecutionOperationalProfile:
        default_profile = self._normalize_profile_name(getattr(settings, "execution_profile_default", PROFILE_STANDARD))
        candidate_profile = default_profile
        source = "env_default"
        source_details: dict[str, Any] = {"origin": "env_default"}
        persisted_overrides: dict[str, int] = {}

        persisted_setting = self.execution_profile_settings.get_active_by_automation_id(automation_id)
        if persisted_setting is not None:
            candidate_profile = self._normalize_profile_name(persisted_setting.execution_profile)
            persisted_overrides = self._extract_persisted_overrides(persisted_setting=persisted_setting)
            source = "persisted_automation"
            source_details = {
                "origin": "persisted_automation",
                "setting_id": str(persisted_setting.id),
                "setting_is_active": bool(persisted_setting.is_active),
                "override_limit_keys": sorted(persisted_overrides.keys()),
            }
        else:
            overrides = getattr(settings, "execution_profile_automation_overrides", {}) or {}
            automation_key = str(automation_id).strip().lower()
            override_profile = self._normalize_profile_name(str(overrides.get(automation_key, "") or ""))
            if override_profile:
                candidate_profile = override_profile
                source = "env_automation_override"
                source_details = {
                    "origin": "env_automation_override",
                    "automation_id": automation_key,
                }

        if candidate_profile not in KNOWN_EXECUTION_PROFILES:
            logger.warning(
                "Invalid execution profile configured; falling back to standard profile.",
                extra={
                    "event": "execution_profile_fallback",
                    "automation_id": str(automation_id),
                    "requested_profile": candidate_profile or "(empty)",
                    "source": source,
                    "fallback_profile": PROFILE_STANDARD,
                },
            )
            candidate_profile = PROFILE_STANDARD
            source = f"{source}_fallback_standard"
            source_details = {
                **source_details,
                "fallback_profile": PROFILE_STANDARD,
            }
            persisted_overrides = {}

        configured_limits = self._profile_limits_from_settings(profile_name=candidate_profile)
        if persisted_overrides:
            configured_limits = self._merge_profile_limits(
                base_limits=configured_limits,
                override_limits=persisted_overrides,
            )
        hard_ceilings = self._hard_limit_ceilings()
        effective_limits, hard_clamped_fields, hard_clamp_details = self._apply_hard_limit_ceilings(
            configured_limits=configured_limits,
            hard_ceilings=hard_ceilings,
        )

        return ExecutionOperationalProfile(
            name=candidate_profile,
            source=source,
            source_details=source_details,
            persisted_overrides=persisted_overrides,
            max_execution_rows=effective_limits[LIMIT_KEY_MAX_EXECUTION_ROWS],
            max_provider_calls=effective_limits[LIMIT_KEY_MAX_PROVIDER_CALLS],
            max_text_chunks=effective_limits[LIMIT_KEY_MAX_TEXT_CHUNKS],
            max_tabular_row_characters=effective_limits[LIMIT_KEY_MAX_TABULAR_ROW_CHARACTERS],
            max_execution_seconds=effective_limits[LIMIT_KEY_MAX_EXECUTION_SECONDS],
            max_context_characters=effective_limits[LIMIT_KEY_MAX_CONTEXT_CHARACTERS],
            max_context_file_characters=effective_limits[LIMIT_KEY_MAX_CONTEXT_FILE_CHARACTERS],
            max_prompt_characters=effective_limits[LIMIT_KEY_MAX_PROMPT_CHARACTERS],
            hard_clamped_fields=hard_clamped_fields,
            hard_clamp_details=hard_clamp_details,
        )

    def build_direct_execution_profile(self, *, source_label: str = "direct_prompt_test") -> ExecutionOperationalProfile:
        candidate_profile = self._normalize_profile_name(getattr(settings, "execution_profile_default", PROFILE_STANDARD))
        if candidate_profile not in KNOWN_EXECUTION_PROFILES:
            logger.warning(
                "Invalid default execution profile configured for direct execution; falling back to standard profile.",
                extra={
                    "event": "execution_profile_fallback",
                    "requested_profile": candidate_profile or "(empty)",
                    "source": source_label,
                    "fallback_profile": PROFILE_STANDARD,
                },
            )
            candidate_profile = PROFILE_STANDARD

        configured_limits = self._profile_limits_from_settings(profile_name=candidate_profile)
        hard_ceilings = self._hard_limit_ceilings()
        effective_limits, hard_clamped_fields, hard_clamp_details = self._apply_hard_limit_ceilings(
            configured_limits=configured_limits,
            hard_ceilings=hard_ceilings,
        )

        return ExecutionOperationalProfile(
            name=candidate_profile,
            source=source_label,
            source_details={"origin": source_label},
            persisted_overrides={},
            max_execution_rows=effective_limits[LIMIT_KEY_MAX_EXECUTION_ROWS],
            max_provider_calls=effective_limits[LIMIT_KEY_MAX_PROVIDER_CALLS],
            max_text_chunks=effective_limits[LIMIT_KEY_MAX_TEXT_CHUNKS],
            max_tabular_row_characters=effective_limits[LIMIT_KEY_MAX_TABULAR_ROW_CHARACTERS],
            max_execution_seconds=effective_limits[LIMIT_KEY_MAX_EXECUTION_SECONDS],
            max_context_characters=effective_limits[LIMIT_KEY_MAX_CONTEXT_CHARACTERS],
            max_context_file_characters=effective_limits[LIMIT_KEY_MAX_CONTEXT_FILE_CHARACTERS],
            max_prompt_characters=effective_limits[LIMIT_KEY_MAX_PROMPT_CHARACTERS],
            hard_clamped_fields=hard_clamped_fields,
            hard_clamp_details=hard_clamp_details,
        )

    def _enforce_execution_rows_profile_limit(
        self,
        *,
        execution_id: UUID,
        total_rows: int,
        execution_profile: ExecutionOperationalProfile,
    ) -> None:
        if execution_profile.is_hard_clamped(LIMIT_KEY_MAX_EXECUTION_ROWS):
            return
        profile_limit = int(execution_profile.max_execution_rows)
        if total_rows <= profile_limit:
            return
        raise AppException(
            "Execution exceeded profile limit of tabular rows.",
            status_code=422,
            code="execution_rows_profile_limit_exceeded",
            details={
                "execution_id": str(execution_id),
                "execution_profile": execution_profile.name,
                "total_rows": int(total_rows),
                "profile_max_execution_rows": profile_limit,
            },
        )

    def _enforce_text_chunks_profile_limit(
        self,
        *,
        execution_id: UUID,
        chunk_count: int,
        execution_profile: ExecutionOperationalProfile,
    ) -> None:
        if execution_profile.is_hard_clamped(LIMIT_KEY_MAX_TEXT_CHUNKS):
            return
        profile_limit = int(execution_profile.max_text_chunks)
        if chunk_count <= profile_limit:
            return
        raise AppException(
            "Execution exceeded profile limit of text chunks.",
            status_code=422,
            code="text_chunks_profile_limit_exceeded",
            details={
                "execution_id": str(execution_id),
                "execution_profile": execution_profile.name,
                "chunk_count": int(chunk_count),
                "profile_max_text_chunks": profile_limit,
            },
        )

    def _enforce_provider_calls_profile_limit(
        self,
        *,
        execution_id: UUID,
        provider_calls: int,
        phase: str,
        execution_profile: ExecutionOperationalProfile,
        row_index: int | None = None,
        chunk_index: int | None = None,
    ) -> None:
        if execution_profile.is_hard_clamped(LIMIT_KEY_MAX_PROVIDER_CALLS):
            return
        profile_limit = int(execution_profile.max_provider_calls)
        next_call_number = provider_calls + 1
        if next_call_number <= profile_limit:
            return
        details: dict[str, Any] = {
            "execution_id": str(execution_id),
            "execution_profile": execution_profile.name,
            "phase": phase,
            "provider_calls": int(provider_calls),
            "next_call_number": int(next_call_number),
            "profile_max_provider_calls": profile_limit,
        }
        if row_index is not None:
            details["row_index"] = int(row_index)
        if chunk_index is not None:
            details["chunk_index"] = int(chunk_index)
        raise AppException(
            "Execution exceeded profile limit of provider calls.",
            status_code=422,
            code="provider_calls_profile_limit_exceeded",
            details=details,
        )

    def _enforce_tabular_row_size_profile_limit(
        self,
        *,
        execution_id: UUID,
        row_index: int,
        row_values: dict[str, Any],
        execution_profile: ExecutionOperationalProfile,
    ) -> None:
        if execution_profile.is_hard_clamped(LIMIT_KEY_MAX_TABULAR_ROW_CHARACTERS):
            return
        profile_limit = int(execution_profile.max_tabular_row_characters)
        row_characters = self._tabular_row_characters(row_values)
        if row_characters <= profile_limit:
            return
        raise AppException(
            "Tabular row exceeded profile character limit.",
            status_code=422,
            code="tabular_row_size_profile_limit_exceeded",
            details={
                "execution_id": str(execution_id),
                "execution_profile": execution_profile.name,
                "row_index": int(row_index),
                "row_characters": int(row_characters),
                "profile_max_tabular_row_characters": profile_limit,
            },
        )

    def _enforce_execution_time_profile_limit(
        self,
        *,
        execution_id: UUID,
        execution_started_at: float,
        phase: str,
        execution_profile: ExecutionOperationalProfile,
        row_index: int | None = None,
        chunk_index: int | None = None,
    ) -> None:
        if execution_profile.is_hard_clamped(LIMIT_KEY_MAX_EXECUTION_SECONDS):
            return
        profile_limit_seconds = int(execution_profile.max_execution_seconds)
        elapsed_seconds = perf_counter() - execution_started_at
        if elapsed_seconds <= profile_limit_seconds:
            return
        details: dict[str, Any] = {
            "execution_id": str(execution_id),
            "execution_profile": execution_profile.name,
            "phase": phase,
            "elapsed_seconds": round(float(elapsed_seconds), 4),
            "profile_max_execution_seconds": profile_limit_seconds,
        }
        if row_index is not None:
            details["row_index"] = int(row_index)
        if chunk_index is not None:
            details["chunk_index"] = int(chunk_index)
        raise AppException(
            "Execution exceeded profile processing time limit.",
            status_code=422,
            code="execution_time_profile_limit_exceeded",
            details=details,
        )

    def _enforce_execution_rows_hard_limit(
        self,
        *,
        execution_id: UUID,
        total_rows: int,
    ) -> None:
        hard_limit = self._safe_hard_limit(settings.max_execution_rows_hard_limit, fallback=1)
        if total_rows <= hard_limit:
            return
        raise AppException(
            "Execution exceeded hard limit of tabular rows.",
            status_code=422,
            code="execution_rows_hard_limit_exceeded",
            details={
                "execution_id": str(execution_id),
                "total_rows": int(total_rows),
                "max_execution_rows_hard_limit": hard_limit,
            },
        )

    def _enforce_text_chunks_hard_limit(
        self,
        *,
        execution_id: UUID,
        chunk_count: int,
    ) -> None:
        hard_limit = self._safe_hard_limit(settings.max_text_chunks_hard_limit, fallback=1)
        if chunk_count <= hard_limit:
            return
        raise AppException(
            "Execution exceeded hard limit of text chunks.",
            status_code=422,
            code="text_chunks_hard_limit_exceeded",
            details={
                "execution_id": str(execution_id),
                "chunk_count": int(chunk_count),
                "max_text_chunks_hard_limit": hard_limit,
            },
        )

    def _enforce_provider_calls_hard_limit(
        self,
        *,
        execution_id: UUID,
        provider_calls: int,
        phase: str,
        row_index: int | None = None,
        chunk_index: int | None = None,
    ) -> None:
        hard_limit = self._safe_hard_limit(settings.max_provider_calls_hard_limit, fallback=1)
        next_call_number = provider_calls + 1
        if next_call_number <= hard_limit:
            return
        details: dict[str, Any] = {
            "execution_id": str(execution_id),
            "phase": phase,
            "provider_calls": int(provider_calls),
            "next_call_number": int(next_call_number),
            "max_provider_calls_hard_limit": hard_limit,
        }
        if row_index is not None:
            details["row_index"] = int(row_index)
        if chunk_index is not None:
            details["chunk_index"] = int(chunk_index)
        raise AppException(
            "Execution exceeded hard limit of provider calls.",
            status_code=422,
            code="provider_calls_hard_limit_exceeded",
            details=details,
        )

    @staticmethod
    def _tabular_row_characters(row_values: dict[str, Any]) -> int:
        return sum(len(str(value or "")) for value in row_values.values())

    def _enforce_tabular_row_size_hard_limit(
        self,
        *,
        execution_id: UUID,
        row_index: int,
        row_values: dict[str, Any],
    ) -> None:
        hard_limit = self._safe_hard_limit(settings.max_tabular_row_characters_hard_limit, fallback=1)
        row_characters = self._tabular_row_characters(row_values)
        if row_characters <= hard_limit:
            return
        raise AppException(
            "Tabular row exceeded hard character limit.",
            status_code=422,
            code="tabular_row_size_hard_limit_exceeded",
            details={
                "execution_id": str(execution_id),
                "row_index": int(row_index),
                "row_characters": int(row_characters),
                "max_tabular_row_characters_hard_limit": hard_limit,
            },
        )

    def _enforce_execution_time_hard_limit(
        self,
        *,
        execution_id: UUID,
        execution_started_at: float,
        phase: str,
        row_index: int | None = None,
        chunk_index: int | None = None,
    ) -> None:
        hard_limit_seconds = self._safe_hard_limit(settings.max_execution_seconds_hard_limit, fallback=1)
        elapsed_seconds = perf_counter() - execution_started_at
        if elapsed_seconds <= hard_limit_seconds:
            return
        details: dict[str, Any] = {
            "execution_id": str(execution_id),
            "phase": phase,
            "elapsed_seconds": round(float(elapsed_seconds), 4),
            "max_execution_seconds_hard_limit": hard_limit_seconds,
        }
        if row_index is not None:
            details["row_index"] = int(row_index)
        if chunk_index is not None:
            details["chunk_index"] = int(chunk_index)
        raise AppException(
            "Execution exceeded hard processing time limit.",
            status_code=422,
            code="execution_time_hard_limit_exceeded",
            details=details,
        )

    def create_execution(
        self,
        *,
        analysis_request_id: UUID,
        request_file_id: UUID | None = None,
        request_file_ids: list[UUID] | None = None,
        input_files: list[Any] | None = None,
        prompt_override: str | None = None,
        api_token: DjangoAiApiToken,
        token_permissions: list[DjangoAiApiTokenPermission],
        ip_address: str | None = None,
        correlation_id: str | None = None,
    ) -> ExecutionCreateResult:
        analysis_request = self.shared_analysis.get_request_by_id(analysis_request_id)
        if analysis_request is None:
            raise AppException(
                "analysis_request_id not found in shared system.",
                status_code=404,
                code="analysis_request_not_found",
                details={"analysis_request_id": str(analysis_request_id)},
            )

        allowed = check_token_permission(
            permissions=token_permissions,
            operation="execution",
            automation_id=analysis_request.automation_id,
        )
        if not allowed:
            raise AppException(
                "Token does not allow execution for this automation.",
                status_code=403,
                code="execution_permission_denied",
            )
        if request_file_ids and input_files:
            raise AppException(
                "Use either request_file_ids or input_files, not both at the same time.",
                status_code=422,
                code="execution_input_payload_conflict",
            )

        resolved_inputs = self._resolve_execution_inputs(
            analysis_request_id=analysis_request_id,
            request_file_id=request_file_id,
            request_file_ids=request_file_ids,
            input_files=input_files,
        )
        self._log_execution_phase(
            phase="execution_create.input_resolution",
            message="Execution create payload resolved to concrete inputs.",
            analysis_request_id=str(analysis_request_id),
            input_summary={
                "input_file_count": len(resolved_inputs),
                "roles": {
                    INPUT_ROLE_PRIMARY: sum(1 for item in resolved_inputs if item.role == INPUT_ROLE_PRIMARY),
                    INPUT_ROLE_CONTEXT: sum(1 for item in resolved_inputs if item.role == INPUT_ROLE_CONTEXT),
                },
                "inputs": [
                    {
                        "request_file_id": str(item.request_file_id),
                        "role": item.role,
                        "order_index": item.order_index,
                    }
                    for item in resolved_inputs
                ],
            },
        )
        primary_input = next((item for item in resolved_inputs if item.role == INPUT_ROLE_PRIMARY), None)
        if primary_input is None:
            raise AppException(
                "Execution input payload does not define a primary file.",
                status_code=422,
                code="execution_primary_input_missing",
            )
        normalized_prompt_override = self._normalize_prompt_override(prompt_override)

        execution = self.shared_executions.create(
            analysis_request_id=analysis_request_id,
            status=ExecutionStatus.PENDING.value,
        )
        self.shared_session.commit()

        queue_job = DjangoAiQueueJob(
            execution_id=execution.id,
            request_file_id=primary_input.request_file_id,
            job_status=ExecutionStatus.QUEUED.value,
            retry_count=0,
            prompt_override_text=normalized_prompt_override,
        )
        self.queue_jobs.add(queue_job)
        for selection in resolved_inputs:
            self.execution_inputs.add(
                DjangoAiExecutionInputFile(
                    execution_id=execution.id,
                    request_file_id=selection.request_file_id,
                    role=selection.role,
                    order_index=selection.order_index,
                )
            )

        serialized_inputs = [
            {
                "request_file_id": str(selection.request_file_id),
                "role": selection.role,
                "order_index": selection.order_index,
            }
            for selection in resolved_inputs
        ]
        self.audit_logs.add(
            DjangoAiAuditLog(
                action_type="execution_created",
                entity_type="analysis_executions",
                entity_id=str(execution.id),
                performed_by_user_id=None,
                changes_json={
                    "analysis_request_id": str(analysis_request_id),
                    "request_file_id": str(primary_input.request_file_id),
                    "request_file_ids": [entry["request_file_id"] for entry in serialized_inputs],
                    "input_files": serialized_inputs,
                    "prompt_override_applied": bool(normalized_prompt_override),
                    "prompt_override_characters": len(normalized_prompt_override or ""),
                    "token_id": str(api_token.id),
                    "queue_job_id": str(queue_job.id),
                },
                ip_address=ip_address,
            )
        )
        self.operational_session.commit()

        self.shared_executions.update_status(execution_id=execution.id, status=ExecutionStatus.QUEUED.value)
        self.shared_session.commit()

        try:
            enqueue_execution_job(
                execution_id=execution.id,
                queue_job_id=queue_job.id,
                correlation_id=correlation_id,
            )
        except Exception as exc:
            logger.exception("Failed to enqueue execution job.", extra={"execution_id": str(execution.id)}, exc_info=exc)
            enqueue_diagnostic = classify_execution_error(
                AppException(
                    "Failed to enqueue execution job.",
                    status_code=500,
                    code="queue_enqueue_failed",
                ),
                failure_phase="execution_create.queue_enqueue",
            )
            self._mark_execution_failed(
                execution_id=execution.id,
                queue_job_id=queue_job.id,
                error_message="Failed to enqueue execution job.",
                worker_name="api",
                ip_address=ip_address,
                register_error_file=False,
                error_diagnostic=enqueue_diagnostic,
            )
            raise AppException(
                "Failed to enqueue execution job.",
                status_code=500,
                code="queue_enqueue_failed",
            ) from exc

        logger.info(
            "Execution created and queued.",
            extra={
                "execution_id": str(execution.id),
                "analysis_request_id": str(analysis_request_id),
                "request_file_id": str(primary_input.request_file_id),
                "request_file_ids": [str(item.request_file_id) for item in resolved_inputs],
                "queue_job_id": str(queue_job.id),
                "prompt_override_applied": bool(normalized_prompt_override),
                "phase": "execution_create.enqueued",
                "event": "execution_create_completed",
            },
        )
        return ExecutionCreateResult(
            execution_id=execution.id,
            queue_job_id=queue_job.id,
            status=ExecutionStatus.QUEUED,
        )

    def get_execution_status(
        self,
        *,
        execution_id: UUID,
        token_permissions: list[DjangoAiApiTokenPermission],
    ) -> ExecutionStatusResult:
        execution = self.shared_executions.get_by_id(execution_id)
        if execution is None:
            raise AppException("Execution not found.", status_code=404, code="execution_not_found")

        analysis_request = self.shared_analysis.get_request_by_id(execution.analysis_request_id)
        if analysis_request is None:
            raise AppException("Related analysis request not found.", status_code=404, code="analysis_request_not_found")

        allowed = check_token_permission(
            permissions=token_permissions,
            operation="execution",
            automation_id=analysis_request.automation_id,
        )
        if not allowed:
            raise AppException("Token cannot access this execution.", status_code=403, code="execution_permission_denied")

        latest_job = self.queue_jobs.get_latest_by_execution_id(execution.id)
        return ExecutionStatusResult(
            execution_id=execution.id,
            status=self._parse_status(execution.status),
            progress=None,
            started_at=latest_job.started_at if latest_job else None,
            finished_at=latest_job.finished_at if latest_job else None,
            error_message=latest_job.error_message if latest_job else None,
            created_at=execution.created_at,
        )

    def list_executions_for_analysis_request(
        self,
        *,
        analysis_request_id: UUID,
        token_permissions: list[DjangoAiApiTokenPermission],
    ) -> list[ExecutionStatusResult]:
        analysis_request = self.shared_analysis.get_request_by_id(analysis_request_id)
        if analysis_request is None:
            raise AppException("analysis_request_id not found.", status_code=404, code="analysis_request_not_found")

        allowed = check_token_permission(
            permissions=token_permissions,
            operation="execution",
            automation_id=analysis_request.automation_id,
        )
        if not allowed:
            raise AppException("Token cannot list executions for this request.", status_code=403, code="execution_permission_denied")

        executions = self.shared_executions.list_by_analysis_request_id(analysis_request_id)
        items: list[ExecutionStatusResult] = []
        for execution in executions:
            latest_job = self.queue_jobs.get_latest_by_execution_id(execution.id)
            items.append(
                ExecutionStatusResult(
                    execution_id=execution.id,
                    status=self._parse_status(execution.status),
                    progress=None,
                    started_at=latest_job.started_at if latest_job else None,
                    finished_at=latest_job.finished_at if latest_job else None,
                    error_message=latest_job.error_message if latest_job else None,
                    created_at=execution.created_at,
                )
            )
        return items

    def list_execution_inputs(
        self,
        *,
        execution_id: UUID,
        token_permissions: list[DjangoAiApiTokenPermission],
    ) -> list[ExecutionInputResult]:
        execution = self.shared_executions.get_by_id(execution_id)
        if execution is None:
            raise AppException("Execution not found.", status_code=404, code="execution_not_found")

        analysis_request = self.shared_analysis.get_request_by_id(execution.analysis_request_id)
        if analysis_request is None:
            raise AppException("Related analysis request not found.", status_code=404, code="analysis_request_not_found")

        allowed = check_token_permission(
            permissions=token_permissions,
            operation="execution",
            automation_id=analysis_request.automation_id,
        )
        if not allowed:
            raise AppException("Token cannot access this execution.", status_code=403, code="execution_permission_denied")

        linked_inputs = self.execution_inputs.list_by_execution_id(execution_id)
        if linked_inputs:
            items: list[ExecutionInputResult] = []
            for linked in linked_inputs:
                request_file = self.request_files.get_by_id(linked.request_file_id)
                items.append(
                    ExecutionInputResult(
                        request_file_id=linked.request_file_id,
                        file_name=request_file.file_name if request_file is not None else None,
                        role=linked.role,
                        order_index=int(linked.order_index or 0),
                        source="linked",
                    )
                )
            return items

        latest_job = self.queue_jobs.get_latest_by_execution_id(execution_id)
        if latest_job is None or latest_job.request_file_id is None:
            return []

        request_file = self.request_files.get_by_id(latest_job.request_file_id)
        return [
            ExecutionInputResult(
                request_file_id=latest_job.request_file_id,
                file_name=request_file.file_name if request_file is not None else None,
                role=INPUT_ROLE_PRIMARY,
                order_index=0,
                source="legacy_queue_job",
            )
        ]

    def _resolve_execution_inputs(
        self,
        *,
        analysis_request_id: UUID,
        request_file_id: UUID | None,
        request_file_ids: list[UUID] | None,
        input_files: list[Any] | None,
    ) -> list[ExecutionInputSelection]:
        raw_entries: list[dict[str, Any]] = []

        if input_files:
            for position, raw in enumerate(input_files):
                if isinstance(raw, dict):
                    file_id = raw.get("request_file_id")
                    role = raw.get("role")
                    order_index = raw.get("order_index")
                else:
                    file_id = getattr(raw, "request_file_id", None)
                    role = getattr(raw, "role", None)
                    order_index = getattr(raw, "order_index", None)
                raw_entries.append(
                    {
                        "request_file_id": file_id,
                        "role": role,
                        "order_index": order_index if order_index is not None else position,
                        "position": position,
                    }
                )
        elif request_file_ids:
            for position, file_id in enumerate(request_file_ids):
                raw_entries.append(
                    {
                        "request_file_id": file_id,
                        "role": None,
                        "order_index": position,
                        "position": position,
                    }
                )
        elif request_file_id is not None:
            raw_entries.append(
                {
                    "request_file_id": request_file_id,
                    "role": INPUT_ROLE_PRIMARY,
                    "order_index": 0,
                    "position": 0,
                }
            )

        if request_file_id is not None and not any(entry["request_file_id"] == request_file_id for entry in raw_entries):
            shifted_entries = []
            for position, entry in enumerate(raw_entries, start=1):
                shifted = dict(entry)
                shifted["order_index"] = position
                shifted["position"] = position
                shifted_entries.append(shifted)
            raw_entries = [
                {
                    "request_file_id": request_file_id,
                    "role": INPUT_ROLE_PRIMARY,
                    "order_index": 0,
                    "position": 0,
                },
                *shifted_entries,
            ]

        if not raw_entries:
            raise AppException(
                "At least one input file is required to create an execution.",
                status_code=422,
                code="execution_input_missing",
            )

        normalized_entries: list[dict[str, Any]] = []
        seen_file_ids: set[UUID] = set()
        for entry in raw_entries:
            raw_file_id = entry.get("request_file_id")
            if raw_file_id is None:
                raise AppException(
                    "Input file entry is missing request_file_id.",
                    status_code=422,
                    code="execution_input_item_invalid",
                )
            try:
                file_id = UUID(str(raw_file_id))
            except ValueError as exc:
                raise AppException(
                    "request_file_id must be a valid UUID.",
                    status_code=422,
                    code="execution_input_item_invalid",
                ) from exc
            if file_id in seen_file_ids:
                raise AppException(
                    "Duplicate request_file_id found in execution input payload.",
                    status_code=422,
                    code="execution_input_duplicate_file",
                    details={"request_file_id": str(file_id)},
                )
            seen_file_ids.add(file_id)

            raw_order_index = entry.get("order_index")
            try:
                order_index = int(raw_order_index)
            except (TypeError, ValueError) as exc:
                raise AppException(
                    "order_index must be an integer value.",
                    status_code=422,
                    code="execution_input_order_invalid",
                ) from exc
            if order_index < 0:
                raise AppException(
                    "order_index must be a non-negative integer.",
                    status_code=422,
                    code="execution_input_order_invalid",
                )

            raw_role = entry.get("role")
            role: str | None = None
            if raw_role is not None:
                role = str(raw_role).strip().lower()
                if role not in ALLOWED_INPUT_ROLES:
                    raise AppException(
                        "Invalid input file role.",
                        status_code=422,
                        code="execution_input_role_invalid",
                        details={"allowed_roles": sorted(ALLOWED_INPUT_ROLES)},
                    )

            normalized_entries.append(
                {
                    "request_file_id": file_id,
                    "role": role,
                    "order_index": order_index,
                    "position": int(entry.get("position", 0)),
                }
            )

        sorted_entries = sorted(
            normalized_entries,
            key=lambda item: (int(item["order_index"]), int(item["position"])),
        )

        explicit_primary_ids = [item["request_file_id"] for item in sorted_entries if item.get("role") == INPUT_ROLE_PRIMARY]
        if len(explicit_primary_ids) > 1:
            raise AppException(
                "Only one input file can be marked as primary.",
                status_code=422,
                code="execution_input_multiple_primary",
            )
        if request_file_id is not None and explicit_primary_ids and explicit_primary_ids[0] != request_file_id:
            raise AppException(
                "request_file_id must match the primary input file when explicit roles are informed.",
                status_code=422,
                code="execution_input_primary_conflict",
            )

        if request_file_id is not None:
            primary_file_id = request_file_id
        elif explicit_primary_ids:
            primary_file_id = explicit_primary_ids[0]
        else:
            primary_file_id = sorted_entries[0]["request_file_id"]

        resolved: list[ExecutionInputSelection] = []
        for index, item in enumerate(sorted_entries):
            current_file_id = item["request_file_id"]
            current_role = INPUT_ROLE_PRIMARY if current_file_id == primary_file_id else INPUT_ROLE_CONTEXT
            request_file = self.request_files.get_by_id(current_file_id)
            if request_file is None:
                raise AppException(
                    "Request file not found.",
                    status_code=404,
                    code="request_file_not_found",
                    details={"request_file_id": str(current_file_id)},
                )
            if request_file.analysis_request_id != analysis_request_id:
                raise AppException(
                    "request_file_id does not belong to analysis_request_id.",
                    status_code=409,
                    code="request_file_analysis_mismatch",
                    details={"request_file_id": str(current_file_id)},
                )
            resolved.append(
                ExecutionInputSelection(
                    request_file_id=current_file_id,
                    role=current_role,
                    order_index=index,
                    file_name=request_file.file_name,
                )
            )
        return resolved

    def process_execution_job(
        self,
        *,
        execution_id: UUID,
        queue_job_id: UUID,
        worker_name: str,
        correlation_id: str | None = None,
    ) -> None:
        failure_phase = "execution.process.initial_validation"
        plan_summary: dict[str, Any] | None = None
        input_summary: dict[str, Any] | None = None
        execution_profile: ExecutionOperationalProfile | None = None
        automation_id: UUID | None = None

        queue_job = self.queue_jobs.get_by_id(queue_job_id)
        if queue_job is None:
            raise AppException("Queue job not found.", status_code=404, code="queue_job_not_found")
        if queue_job.execution_id != execution_id:
            raise AppException(
                "queue_job_id does not match execution_id.",
                status_code=409,
                code="queue_job_execution_mismatch",
            )

        shared_execution = self.shared_executions.get_by_id(execution_id)
        if shared_execution is None:
            raise AppException("Execution not found.", status_code=404, code="execution_not_found")

        if shared_execution.status == ExecutionStatus.COMPLETED.value:
            logger.info("Execution already completed. Skipping duplicate worker run.", extra={"execution_id": str(execution_id)})
            return

        if self._is_concurrency_limited(queue_job_id=queue_job_id):
            concurrency_diagnostic = ExecutionErrorDiagnostic(
                message="Global concurrency limit reached.",
                error_code="concurrency_limit_reached",
                error_category="orchestration",
                failure_phase="execution.process.concurrency_gate",
            )
            self._schedule_retry(
                execution_id=execution_id,
                queue_job=queue_job,
                reason="Global concurrency limit reached.",
                worker_name=worker_name,
                correlation_id=correlation_id,
                error_diagnostic=concurrency_diagnostic,
            )
            return

        acquired = self.queue_jobs.acquire_for_processing(
            queue_job_id=queue_job_id,
            worker_name=worker_name,
            started_at=datetime.now(timezone.utc),
        )
        self.operational_session.commit()
        if not acquired:
            logger.info(
                "Queue job was already acquired by another worker. Skipping duplicate run.",
                extra={"execution_id": str(execution_id), "queue_job_id": str(queue_job_id)},
            )
            return

        queue_job = self.queue_jobs.get_by_id(queue_job_id)
        if queue_job is None:
            return

        execution_started_at = perf_counter()
        try:
            self._log_execution_phase(
                phase="execution.process.start",
                message="Execution worker run started.",
                execution_id=str(execution_id),
                queue_job_id=str(queue_job_id),
                worker_name=worker_name,
            )
            self.shared_executions.update_status(execution_id=execution_id, status=ExecutionStatus.PROCESSING.value)
            self.shared_session.commit()
            self.audit_logs.add(
                DjangoAiAuditLog(
                    action_type="execution_started",
                    entity_type="analysis_executions",
                    entity_id=str(execution_id),
                    performed_by_user_id=None,
                    changes_json={"queue_job_id": str(queue_job_id), "worker_name": worker_name},
                    ip_address=None,
                )
            )
            self.operational_session.commit()
            logger.info("Execution processing started.", extra={"execution_id": str(execution_id), "worker_name": worker_name})

            failure_phase = "execution.process.input_resolution"
            self._log_execution_phase(
                phase=failure_phase,
                message="Resolving execution inputs for worker run.",
                execution_id=str(execution_id),
                queue_job_id=str(queue_job_id),
            )
            processing_inputs = self._load_execution_processing_inputs(
                execution_id=execution_id,
                queue_job=queue_job,
            )
            input_summary = summarize_processing_inputs(processing_inputs)
            self._log_execution_phase(
                phase=f"{failure_phase}.completed",
                message="Execution inputs resolved.",
                execution_id=str(execution_id),
                queue_job_id=str(queue_job_id),
                **input_summary,
            )

            failure_phase = "execution.process.strategy_resolution"
            processing_plan = self._resolve_processing_strategy(processing_inputs=processing_inputs)
            plan_summary = summarize_processing_plan(processing_plan)
            if queue_job.request_file_id != processing_plan.primary_input.request_file_id:
                queue_job.request_file_id = processing_plan.primary_input.request_file_id
                self.operational_session.commit()
            self._log_execution_phase(
                phase=f"{failure_phase}.completed",
                message="Execution strategy resolved.",
                execution_id=str(execution_id),
                queue_job_id=str(queue_job_id),
                **plan_summary,
            )

            failure_phase = "execution.process.runtime_resolution"
            shared_request = self.shared_analysis.get_request_by_id(shared_execution.analysis_request_id)
            if shared_request is None:
                raise AppException("Related analysis request not found.", status_code=404, code="analysis_request_not_found")
            automation_id = shared_request.automation_id

            failure_phase = "execution.process.profile_resolution"
            profile_automation_id = shared_request.automation_id
            execution_profile = self._resolve_execution_profile(automation_id=profile_automation_id)
            self._log_execution_phase(
                phase=f"{failure_phase}.completed",
                message="Execution profile resolved.",
                execution_id=str(execution_id),
                queue_job_id=str(queue_job_id),
                automation_id=str(profile_automation_id),
                execution_profile=execution_profile.name,
                execution_profile_source=execution_profile.source,
                execution_profile_source_details=execution_profile.source_details,
                execution_profile_limits=execution_profile.to_limits_dict(),
                execution_profile_persisted_overrides=execution_profile.persisted_overrides,
                hard_clamped_fields=list(execution_profile.hard_clamped_fields),
                hard_clamp_details=execution_profile.hard_clamp_details,
            )

            failure_phase = "execution.process.runtime_resolution"
            prompt_override = self._normalize_prompt_override(queue_job.prompt_override_text)
            resolved_runtime = self.runtime_resolver.resolve(
                shared_request.automation_id,
                require_prompt=not bool(prompt_override),
            )
            effective_prompt = prompt_override or resolved_runtime.prompt_text
            debug_enabled = bool(getattr(resolved_runtime, "debug_enabled", False))
            logger.info(
                "Execution runtime resolved from shared system.",
                extra={
                    "execution_id": str(execution_id),
                    "event": "runtime_config_resolved",
                    "provider": resolved_runtime.provider_slug,
                    "model": resolved_runtime.model_slug,
                    "prompt_version": resolved_runtime.prompt_version,
                    "prompt_source": "override" if prompt_override else "official",
                    "debug_enabled": debug_enabled,
                },
            )
            if prompt_override:
                self._log_execution_phase(
                    phase="execution.process.prompt_override",
                    message="Prompt override detected for this execution; official prompt remains unchanged.",
                    execution_id=str(execution_id),
                    queue_job_id=str(queue_job_id),
                    prompt_override_characters=len(prompt_override),
                )

            runtime = self.provider_service.resolve_runtime(
                provider_slug=resolved_runtime.provider_slug,
                model_slug=resolved_runtime.model_slug,
                credential_id=getattr(resolved_runtime, "credential_id", None),
            )
            logger.info(
                "Operational provider/model validation succeeded.",
                extra={
                    "execution_id": str(execution_id),
                    "event": "runtime_validation_ok",
                    "provider": runtime.provider.slug,
                    "model": runtime.model.model_slug,
                },
            )

            failure_phase = "execution.process.output_contract_resolution"
            runtime_output_type = getattr(resolved_runtime, "output_type", None)
            runtime_result_parser = getattr(resolved_runtime, "result_parser", None)
            runtime_result_formatter = getattr(resolved_runtime, "result_formatter", None)
            runtime_output_schema = getattr(resolved_runtime, "output_schema", None)
            explicit_output_contract_present = any(
                value is not None
                for value in (
                    runtime_output_type,
                    runtime_result_parser,
                    runtime_result_formatter,
                    runtime_output_schema,
                )
            )

            try:
                resolved_output_contract = self._resolve_execution_output_contract(
                    automation_id=shared_request.automation_id,
                    automation_slug=getattr(resolved_runtime, "automation_slug", None),
                    processing_plan=processing_plan,
                    runtime_output_type=runtime_output_type,
                    runtime_result_parser=runtime_result_parser,
                    runtime_result_formatter=runtime_result_formatter,
                    runtime_output_schema=runtime_output_schema,
                )
            except AppException as contract_exc:
                if contract_exc.payload.code in {
                    "execution_output_contract_invalid",
                    "execution_output_schema_invalid",
                    "execution_output_contract_incompatible",
                }:
                    self._log_execution_phase(
                        phase=f"{failure_phase}.invalid",
                        message="Execution output contract is invalid for this automation.",
                        level="error",
                        execution_id=str(execution_id),
                        queue_job_id=str(queue_job_id),
                        contract_configured=explicit_output_contract_present,
                        automation_id=str(shared_request.automation_id),
                        output_type=runtime_output_type,
                        parser_strategy=runtime_result_parser,
                        formatter_strategy=runtime_result_formatter,
                        output_schema_payload_type=type(runtime_output_schema).__name__
                        if runtime_output_schema is not None
                        else None,
                        error_code=contract_exc.payload.code,
                    )
                raise

            processing_plan = self.strategy_engine.with_output_contract(
                processing_plan=processing_plan,
                output_contract=resolved_output_contract,
            )
            plan_summary = summarize_processing_plan(processing_plan)
            if resolved_output_contract.source == "fallback_no_output_contract_config":
                self._log_execution_phase(
                    phase=f"{failure_phase}.fallback",
                    message="Execution output contract fallback was applied because automation has no explicit configuration.",
                    execution_id=str(execution_id),
                    queue_job_id=str(queue_job_id),
                    automation_id=str(shared_request.automation_id),
                    output_contract_source=resolved_output_contract.source,
                )
            self._log_execution_phase(
                phase=f"{failure_phase}.completed",
                message="Execution output contract resolved.",
                execution_id=str(execution_id),
                queue_job_id=str(queue_job_id),
                output_schema_columns=list(processing_plan.output_contract.output_schema.columns),
                **(plan_summary or {}),
            )

            self.shared_executions.update_status(execution_id=execution_id, status=ExecutionStatus.GENERATING_OUTPUT.value)
            self.shared_session.commit()

            failure_phase = "execution.process.pipeline_run"
            self._log_execution_phase(
                phase=failure_phase,
                message="Execution pipeline started.",
                execution_id=str(execution_id),
                queue_job_id=str(queue_job_id),
                execution_profile=execution_profile.name if execution_profile else None,
                **(plan_summary or {}),
            )
            if execution_profile is None:
                raise AppException(
                    "Execution profile could not be resolved.",
                    status_code=500,
                    code="execution_profile_resolution_failed",
                )
            self._enforce_execution_time_profile_limit(
                execution_id=execution_id,
                execution_started_at=execution_started_at,
                phase="execution.process.pipeline_gate",
                execution_profile=execution_profile,
            )
            self._enforce_execution_time_hard_limit(
                execution_id=execution_id,
                execution_started_at=execution_started_at,
                phase="execution.process.pipeline_gate",
            )
            processed_output = self._process_execution_by_strategy(
                execution_id=execution_id,
                processing_plan=processing_plan,
                official_prompt=effective_prompt,
                runtime=runtime,
                execution_started_at=execution_started_at,
                execution_profile=execution_profile,
                debug_enabled=debug_enabled,
                automation_id=automation_id,
                retry_count=max(int(queue_job.retry_count or 0), 0),
            )
            self._log_execution_phase(
                phase=f"{failure_phase}.completed",
                message="Execution pipeline finished.",
                execution_id=str(execution_id),
                queue_job_id=str(queue_job_id),
                provider_calls=processed_output.provider_calls,
                **processed_output.processing_summary,
            )

            failure_phase = "execution.process.output_persist"
            self.file_service.register_generated_execution_file(
                execution_id=execution_id,
                file_type="output",
                file_name=processed_output.file_name,
                content=processed_output.content,
                mime_type=processed_output.mime_type,
            )
            for auxiliary_file in processed_output.auxiliary_files:
                self.file_service.register_generated_execution_file(
                    execution_id=execution_id,
                    file_type=auxiliary_file.file_type,
                    file_name=auxiliary_file.file_name,
                    content=auxiliary_file.content,
                    mime_type=auxiliary_file.mime_type,
                )
            self._enforce_execution_time_profile_limit(
                execution_id=execution_id,
                execution_started_at=execution_started_at,
                phase="execution.process.output_persist",
                execution_profile=execution_profile,
            )
            self._enforce_execution_time_hard_limit(
                execution_id=execution_id,
                execution_started_at=execution_started_at,
                phase="execution.process.output_persist",
            )
            self._log_execution_phase(
                phase=f"{failure_phase}.completed",
                message="Execution output file persisted.",
                execution_id=str(execution_id),
                queue_job_id=str(queue_job_id),
                output_file_name=processed_output.file_name,
                output_file_mime=processed_output.mime_type,
                output_file_size=len(processed_output.content),
                auxiliary_files_count=len(processed_output.auxiliary_files),
            )

            failure_phase = "execution.process.final_persist"
            queue_job.job_status = ExecutionStatus.COMPLETED.value
            queue_job.error_message = None
            queue_job.finished_at = datetime.now(timezone.utc)
            self.shared_executions.update_status(execution_id=execution_id, status=ExecutionStatus.COMPLETED.value)
            self.audit_logs.add(
                DjangoAiAuditLog(
                    action_type="execution_completed",
                    entity_type="analysis_executions",
                    entity_id=str(execution_id),
                    performed_by_user_id=None,
                    changes_json={
                        "queue_job_id": str(queue_job_id),
                        "providers_used": sorted(processed_output.providers_used),
                        "models_used": sorted(processed_output.models_used),
                        "input_tokens": processed_output.total_input_tokens,
                        "output_tokens": processed_output.total_output_tokens,
                        "estimated_cost": str(processed_output.total_cost),
                    "provider_calls": processed_output.provider_calls,
                    "processing_summary": processed_output.processing_summary,
                    "auxiliary_files": [
                        {
                            "file_type": item.file_type,
                            "file_name": item.file_name,
                            "mime_type": item.mime_type,
                            "file_size": len(item.content),
                        }
                        for item in processed_output.auxiliary_files
                    ],
                    "execution_profile": execution_profile.name if execution_profile else None,
                    "execution_profile_source": execution_profile.source if execution_profile else None,
                    "execution_profile_source_details": execution_profile.source_details if execution_profile else None,
                    "execution_profile_limits": execution_profile.to_limits_dict() if execution_profile else None,
                    "execution_profile_persisted_overrides": execution_profile.persisted_overrides
                    if execution_profile
                    else None,
                    "execution_profile_hard_clamped_fields": list(execution_profile.hard_clamped_fields)
                    if execution_profile
                    else None,
                    "prompt_override_applied": bool(prompt_override),
                    "prompt_source": "override" if prompt_override else "official",
                },
                ip_address=None,
            )
        )
            self.operational_session.commit()
            self.shared_session.commit()
            logger.info(
                "Execution processing completed.",
                extra={
                    "execution_id": str(execution_id),
                    "provider": ",".join(sorted(processed_output.providers_used))
                    if processed_output.providers_used
                    else None,
                    "model": ",".join(sorted(processed_output.models_used))
                    if processed_output.models_used
                    else None,
                    "input_tokens": processed_output.total_input_tokens,
                    "output_tokens": processed_output.total_output_tokens,
                    "estimated_cost": str(processed_output.total_cost),
                    "provider_calls": processed_output.provider_calls,
                    "execution_profile": execution_profile.name if execution_profile else None,
                    "phase": "execution.process.completed",
                    "duration_seconds": round(perf_counter() - execution_started_at, 4),
                },
            )
        except Exception as exc:
            diagnostic = classify_execution_error(exc, failure_phase=failure_phase)
            provider_details = self._provider_error_details(exc)
            provider_status_code = provider_details.get("status_code") or provider_details.get("http_status_code")
            logger.exception(
                "Execution processing failed.",
                extra={
                    "execution_id": str(execution_id),
                    "queue_job_id": str(queue_job_id),
                    "phase": diagnostic.failure_phase,
                    "error_code": diagnostic.error_code,
                    "error_category": diagnostic.error_category,
                    "input_file_count": input_summary.get("input_file_count") if input_summary else None,
                    "input_type": plan_summary.get("input_type") if plan_summary else None,
                    "processing_mode": plan_summary.get("processing_mode") if plan_summary else None,
                    "output_type": plan_summary.get("output_type") if plan_summary else None,
                    "parser_strategy": plan_summary.get("parser_strategy") if plan_summary else None,
                    "execution_profile": execution_profile.name if execution_profile else None,
                    "automation_id": str(automation_id) if automation_id is not None else None,
                    "provider": provider_details.get("provider_slug")
                    or provider_details.get("provider")
                    or None,
                    "model": provider_details.get("model_slug") or None,
                    "status_code": int(provider_status_code)
                    if isinstance(provider_status_code, int) or str(provider_status_code).isdigit()
                    else None,
                    "error_type": self._classify_execution_error_type(
                        exc=exc,
                        stage_of_failure="provider_call",
                    ),
                    "request_id": provider_details.get("provider_request_id") or None,
                    "client_request_id": provider_details.get("client_request_id") or None,
                    "duration_ms": int(provider_details.get("duration_ms"))
                    if isinstance(provider_details.get("duration_ms"), int)
                    or str(provider_details.get("duration_ms")).isdigit()
                    else None,
                },
                exc_info=exc,
            )
            if self._should_retry(exc=exc, retry_count=queue_job.retry_count or 0):
                self._schedule_retry(
                    execution_id=execution_id,
                    queue_job=queue_job,
                    reason=diagnostic.message,
                    worker_name=worker_name,
                    correlation_id=correlation_id,
                    error_diagnostic=diagnostic,
                )
                return

            self._mark_execution_failed(
                execution_id=execution_id,
                queue_job_id=queue_job_id,
                error_message=diagnostic.message,
                worker_name=worker_name,
                ip_address=None,
                register_error_file=True,
                error_diagnostic=diagnostic,
            )

    def _load_execution_processing_inputs(
        self,
        *,
        execution_id: UUID,
        queue_job: DjangoAiQueueJob,
    ) -> list[EngineExecutionInput]:
        linked_inputs = self.execution_inputs.list_by_execution_id(execution_id)
        if linked_inputs:
            resolved: list[EngineExecutionInput] = []
            for linked in linked_inputs:
                request_file = self.request_files.get_by_id(linked.request_file_id)
                if request_file is None:
                    raise AppException(
                        "Request file not found for execution input.",
                        status_code=404,
                        code="request_file_not_found",
                        details={"request_file_id": str(linked.request_file_id)},
                    )
                file_name = str(request_file.file_name or "")
                role = str(linked.role or "").strip().lower()
                resolved.append(
                    EngineExecutionInput(
                        request_file_id=linked.request_file_id,
                        role=role,
                        order_index=int(linked.order_index or 0),
                        file_name=file_name,
                        file_path=str(request_file.file_path or ""),
                        mime_type=request_file.mime_type,
                        file_kind=self.strategy_engine.detect_file_kind(
                            file_name=file_name,
                            mime_type=request_file.mime_type,
                        ),
                        source="linked",
                    )
                )
            return resolved

        request_file_id = queue_job.request_file_id
        if request_file_id is None:
            raise AppException(
                "Queue job is missing request_file_id and no linked execution inputs were found.",
                status_code=409,
                code="queue_job_request_file_missing",
            )

        request_file = self.request_files.get_by_id(request_file_id)
        if request_file is None:
            raise AppException(
                "Request file not found for execution.",
                status_code=404,
                code="request_file_not_found",
                details={"request_file_id": str(request_file_id)},
            )

        file_name = str(request_file.file_name or "")
        return [
            EngineExecutionInput(
                request_file_id=request_file_id,
                role=INPUT_ROLE_PRIMARY,
                order_index=0,
                file_name=file_name,
                file_path=str(request_file.file_path or ""),
                mime_type=request_file.mime_type,
                file_kind=self.strategy_engine.detect_file_kind(
                    file_name=file_name,
                    mime_type=request_file.mime_type,
                ),
                source="legacy_queue_job",
            )
        ]

    def _resolve_processing_strategy(
        self,
        *,
        processing_inputs: list[EngineExecutionInput],
        output_contract: ExecutionOutputContract | None = None,
    ) -> EngineExecutionPlan:
        return self.strategy_engine.resolve_plan(
            processing_inputs=processing_inputs,
            output_contract=output_contract,
        )

    def _resolve_execution_output_contract(
        self,
        *,
        automation_id: UUID | None,
        automation_slug: str | None,
        processing_plan: EngineExecutionPlan,
        runtime_output_type: str | None,
        runtime_result_parser: str | None,
        runtime_result_formatter: str | None,
        runtime_output_schema: dict[str, Any] | str | None,
    ) -> ExecutionOutputContract:
        return self.output_contract_resolver.resolve(
            input_type=processing_plan.input_type,
            automation_id=automation_id,
            automation_slug=automation_slug,
            runtime_output_type=runtime_output_type,
            runtime_result_parser=runtime_result_parser,
            runtime_result_formatter=runtime_result_formatter,
            runtime_output_schema=runtime_output_schema,
        )

    def _process_execution_by_strategy(
        self,
        *,
        execution_id: UUID,
        processing_plan: EngineExecutionPlan,
        official_prompt: str,
        runtime: ProviderRuntimeSelection,
        execution_started_at: float,
        execution_profile: ExecutionOperationalProfile,
        debug_enabled: bool = False,
        automation_id: UUID | None = None,
        retry_count: int = 0,
        progress_callback: Callable[[ExecutionProgressUpdate], None] | None = None,
    ) -> ProcessedOutput:
        self._enforce_execution_time_hard_limit(
            execution_id=execution_id,
            execution_started_at=execution_started_at,
            phase="execution.pipeline.dispatch",
        )
        self._enforce_execution_time_profile_limit(
            execution_id=execution_id,
            execution_started_at=execution_started_at,
            phase="execution.pipeline.dispatch",
            execution_profile=execution_profile,
        )
        if processing_plan.input_type == ExecutionInputType.TABULAR_WITH_CONTEXT:
            global_context = self._build_global_context_text(
                context_inputs=processing_plan.context_inputs,
                execution_profile=execution_profile,
            )
            return self._process_tabular_file(
                execution_id=execution_id,
                file_path=processing_plan.primary_input.file_path,
                file_name=processing_plan.primary_input.file_name,
                official_prompt=official_prompt,
                runtime=runtime,
                global_context=global_context,
                parser_strategy=processing_plan.parser_strategy,
                formatter_strategy=processing_plan.formatter_strategy,
                output_type=processing_plan.output_type,
                output_contract=processing_plan.output_contract,
                execution_started_at=execution_started_at,
                execution_profile=execution_profile,
                debug_enabled=debug_enabled,
                automation_id=automation_id,
                retry_count=retry_count,
                progress_callback=progress_callback,
            )

        if processing_plan.input_type == ExecutionInputType.TABULAR:
            return self._process_tabular_file(
                execution_id=execution_id,
                file_path=processing_plan.primary_input.file_path,
                file_name=processing_plan.primary_input.file_name,
                official_prompt=official_prompt,
                runtime=runtime,
                global_context=None,
                parser_strategy=processing_plan.parser_strategy,
                formatter_strategy=processing_plan.formatter_strategy,
                output_type=processing_plan.output_type,
                output_contract=processing_plan.output_contract,
                execution_started_at=execution_started_at,
                execution_profile=execution_profile,
                debug_enabled=debug_enabled,
                automation_id=automation_id,
                retry_count=retry_count,
                progress_callback=progress_callback,
            )

        if processing_plan.input_type == ExecutionInputType.TEXT:
            return self._process_text_file(
                execution_id=execution_id,
                file_path=processing_plan.primary_input.file_path,
                file_name=processing_plan.primary_input.file_name,
                official_prompt=official_prompt,
                runtime=runtime,
                parser_strategy=processing_plan.parser_strategy,
                formatter_strategy=processing_plan.formatter_strategy,
                output_type=processing_plan.output_type,
                output_contract=processing_plan.output_contract,
                execution_started_at=execution_started_at,
                execution_profile=execution_profile,
                debug_enabled=debug_enabled,
                automation_id=automation_id,
                retry_count=retry_count,
                progress_callback=progress_callback,
            )

        if processing_plan.input_type == ExecutionInputType.MULTI_TEXT:
            merged_content = self._combine_textual_inputs(
                execution_inputs=processing_plan.ordered_inputs,
                execution_profile=execution_profile,
            )
            return self._process_text_content(
                execution_id=execution_id,
                file_content=merged_content,
                official_prompt=official_prompt,
                runtime=runtime,
                parser_strategy=processing_plan.parser_strategy,
                formatter_strategy=processing_plan.formatter_strategy,
                output_type=processing_plan.output_type,
                output_contract=processing_plan.output_contract,
                execution_started_at=execution_started_at,
                execution_profile=execution_profile,
                debug_enabled=debug_enabled,
                automation_id=automation_id,
                retry_count=retry_count,
                progress_callback=progress_callback,
            )

        raise AppException(
            "Execution input strategy is invalid for processing.",
            status_code=422,
            code="invalid_execution_input_combination",
            details={"input_type": processing_plan.input_type.value},
        )

    def _process_text_file(
        self,
        *,
        execution_id: UUID,
        file_path: str,
        file_name: str,
        official_prompt: str,
        runtime: ProviderRuntimeSelection,
        parser_strategy: ExecutionParserStrategy,
        formatter_strategy: ExecutionFormatterStrategy,
        output_type: ExecutionOutputType,
        output_contract: ExecutionOutputContract,
        execution_started_at: float,
        execution_profile: ExecutionOperationalProfile,
        debug_enabled: bool = False,
        automation_id: UUID | None = None,
        retry_count: int = 0,
        progress_callback: Callable[[ExecutionProgressUpdate], None] | None = None,
    ) -> ProcessedOutput:
        self._notify_progress_update(
            progress_callback=progress_callback,
            phase="reading_input",
            status_message="Lendo arquivo textual de entrada.",
        )
        self._log_execution_phase(
            phase="execution.pipeline.file_read",
            message="Reading textual input file.",
            execution_id=str(execution_id),
            file_extension=Path(file_name).suffix.lower(),
            parser_strategy=parser_strategy.value,
            output_type=output_type.value,
        )
        input_file_content = self._read_input_file_content(
            file_path=file_path,
            file_name=file_name,
        )
        self._log_execution_phase(
            phase="execution.pipeline.file_read.completed",
            message="Textual input file loaded.",
            execution_id=str(execution_id),
            input_characters=len(input_file_content or ""),
        )
        self._notify_progress_update(
            progress_callback=progress_callback,
            phase="prompt_build",
            status_message="Entrada textual carregada. Preparando processamento.",
        )
        return self._process_text_content(
            execution_id=execution_id,
            file_content=input_file_content,
            official_prompt=official_prompt,
            runtime=runtime,
            parser_strategy=parser_strategy,
            formatter_strategy=formatter_strategy,
            output_type=output_type,
            output_contract=output_contract,
            execution_started_at=execution_started_at,
            execution_profile=execution_profile,
            debug_enabled=debug_enabled,
            automation_id=automation_id,
            retry_count=retry_count,
            progress_callback=progress_callback,
        )

    def _process_text_content(
        self,
        *,
        execution_id: UUID,
        file_content: str,
        official_prompt: str,
        runtime: ProviderRuntimeSelection,
        parser_strategy: ExecutionParserStrategy,
        formatter_strategy: ExecutionFormatterStrategy,
        output_type: ExecutionOutputType,
        output_contract: ExecutionOutputContract,
        execution_started_at: float,
        execution_profile: ExecutionOperationalProfile,
        debug_enabled: bool = False,
        automation_id: UUID | None = None,
        retry_count: int = 0,
        progress_callback: Callable[[ExecutionProgressUpdate], None] | None = None,
    ) -> ProcessedOutput:
        content_chunks = self._chunk_content(file_content)
        self._enforce_text_chunks_profile_limit(
            execution_id=execution_id,
            chunk_count=len(content_chunks),
            execution_profile=execution_profile,
        )
        self._enforce_text_chunks_hard_limit(
            execution_id=execution_id,
            chunk_count=len(content_chunks),
        )
        provider_calls = 0
        self._log_execution_phase(
            phase="execution.pipeline.prompt_build",
            message="Text pipeline prepared content chunks.",
            execution_id=str(execution_id),
            chunk_count=len(content_chunks),
            parser_strategy=parser_strategy.value,
            output_type=output_type.value,
        )
        self._notify_progress_update(
            progress_callback=progress_callback,
            phase="processing_chunks",
            status_message="Processando conteudo textual no modelo.",
            processed_chunks=0,
            total_chunks=len(content_chunks),
        )

        total_input_tokens = 0
        total_output_tokens = 0
        total_cost = Decimal("0")
        output_chunks: list[str] = []
        providers_used: set[str] = set()
        models_used: set[str] = set()
        debug_chunks: list[dict[str, Any]] = []
        chunk_retry_count = max(int(retry_count or 0), 0)

        for chunk_index, content_chunk in enumerate(content_chunks, start=1):
            self._notify_progress_update(
                progress_callback=progress_callback,
                phase="processing_chunks",
                status_message="Processando conteudo textual no modelo.",
                processed_chunks=chunk_index - 1,
                total_chunks=len(content_chunks),
                current_row=chunk_index,
            )
            self._enforce_execution_time_profile_limit(
                execution_id=execution_id,
                execution_started_at=execution_started_at,
                phase="execution.pipeline.text_chunk_loop",
                execution_profile=execution_profile,
                chunk_index=chunk_index,
            )
            self._enforce_execution_time_hard_limit(
                execution_id=execution_id,
                execution_started_at=execution_started_at,
                phase="execution.pipeline.text_chunk_loop",
                chunk_index=chunk_index,
            )
            prompt_input = self._build_provider_prompt(
                official_prompt=official_prompt,
                file_content=content_chunk,
                execution_profile=execution_profile,
            )
            chunk_stage_of_failure = "provider_call"
            chunk_debug: dict[str, Any] | None = None
            if debug_enabled:
                chunk_debug = {
                    "chunk_index": chunk_index,
                    "chunk_characters": len(content_chunk or ""),
                    "prompt_template": official_prompt,
                    "prompt_final": "",
                    "prompt_truncated": False,
                    **self._build_provider_debug_context(
                        runtime=runtime,
                        retry_count=chunk_retry_count,
                    ),
                    "provider_parameters": {
                        "temperature": getattr(settings, "temperature", None),
                        "max_tokens": getattr(settings, "max_tokens", None),
                    },
                    "response_raw_text": "",
                    "response_raw_payload": None,
                    "parsed_output": None,
                    "normalized_output": "",
                    "warnings": [],
                    "errors": [],
                }
            self._enforce_provider_calls_profile_limit(
                execution_id=execution_id,
                provider_calls=provider_calls,
                phase="execution.pipeline.provider_call",
                execution_profile=execution_profile,
                chunk_index=chunk_index,
            )
            self._enforce_provider_calls_hard_limit(
                execution_id=execution_id,
                provider_calls=provider_calls,
                phase="execution.pipeline.provider_call",
                chunk_index=chunk_index,
            )
            self._log_execution_phase(
                phase="execution.pipeline.provider_call",
                message="Executing provider call for text chunk.",
                level="debug",
                execution_id=str(execution_id),
                chunk_index=chunk_index,
                chunk_characters=len(content_chunk),
            )
            sanitized_prompt, was_truncated = self._enforce_token_limit(
                prompt=prompt_input,
                provider_runtime=runtime,
            )
            if chunk_debug is not None:
                chunk_debug["prompt_final"] = sanitized_prompt
                chunk_debug["prompt_truncated"] = bool(was_truncated)
                chunk_payload = self._build_provider_request_payload(
                    runtime=runtime,
                    prompt_input=sanitized_prompt,
                )
                chunk_debug["request_payload_sanitized"] = chunk_payload
                chunk_token_param = ""
                if isinstance(chunk_payload, dict):
                    if "max_completion_tokens" in chunk_payload:
                        chunk_token_param = "max_completion_tokens"
                    elif "max_tokens" in chunk_payload:
                        chunk_token_param = "max_tokens"
                chunk_debug["token_limit_param_used"] = chunk_token_param
                chunk_debug["api_family_resolved"] = "chat_completions"
                if str(runtime.provider.slug or "").strip().lower() == "openai":
                    chunk_debug["request_profile_resolved"] = (
                        "gpt5_chat" if chunk_token_param == "max_completion_tokens" else "legacy_chat"
                    )
            if was_truncated:
                logger.warning(
                    "Prompt content truncated due to token limit.",
                    extra={
                        "execution_id": str(execution_id),
                        "event": "content_truncated",
                        "chunk_index": chunk_index,
                    },
                )
                if chunk_debug is not None:
                    chunk_debug["warnings"] = [*chunk_debug["warnings"], "prompt_truncated_to_token_limit"]

            chunk_client_request_id = f"{execution_id}:chunk:{chunk_index}:{uuid4().hex[:8]}"
            provider_call_started_at = datetime.now(timezone.utc).isoformat()
            provider_call_started_perf = perf_counter()
            if chunk_debug is not None:
                chunk_debug["started_at"] = provider_call_started_at
                chunk_debug["stage_of_failure"] = chunk_stage_of_failure
                chunk_debug["client_request_id"] = chunk_client_request_id
            try:
                provider_result = self._execute_with_runtime(
                    prompt_input=sanitized_prompt,
                    runtime=runtime,
                    client_request_id=chunk_client_request_id,
                )
            except Exception as chunk_exc:
                if chunk_debug is not None:
                    self._enrich_debug_with_error(
                        debug_item=chunk_debug,
                        exc=chunk_exc,
                        stage_of_failure=chunk_stage_of_failure,
                    )
                    chunk_debug["errors"] = [self._summarize_execution_error(chunk_exc)]
                    debug_chunks.append(chunk_debug)
                if self._classify_execution_error_type(
                    exc=chunk_exc,
                    stage_of_failure=chunk_stage_of_failure,
                ).startswith("provider_"):
                    self._log_provider_error(
                        execution_id=execution_id,
                        row_index=None,
                        automation_id=automation_id,
                        runtime=runtime,
                        stage_of_failure=chunk_stage_of_failure,
                        exc=chunk_exc,
                    )
                raise
            provider_call_finished_at = datetime.now(timezone.utc).isoformat()
            provider_call_duration_ms = max(int((perf_counter() - provider_call_started_perf) * 1000), 0)
            if chunk_debug is not None:
                chunk_debug["finished_at"] = provider_call_finished_at
                chunk_debug["duration_ms"] = provider_call_duration_ms
                chunk_debug["http_status_code"] = 200
                chunk_debug["stage_of_failure"] = "provider_response_validation"
                chunk_debug["response_raw_text"] = str(provider_result.output_text or "")
                chunk_debug["response_raw_payload"] = provider_result.raw_response
            provider_calls += 1
            chunk_cost = runtime.client.estimate_cost(
                input_tokens=provider_result.input_tokens,
                output_tokens=provider_result.output_tokens,
                cost_input_per_1k_tokens=runtime.model.cost_input_per_1k_tokens,
                cost_output_per_1k_tokens=runtime.model.cost_output_per_1k_tokens,
            )
            if total_cost + chunk_cost > Decimal(str(settings.max_cost_per_execution)):
                raise AppException(
                    "Execution aborted due to estimated cost limit.",
                    status_code=422,
                    code="cost_limit_exceeded",
                    details={"max_cost_per_execution": settings.max_cost_per_execution},
                )

            self.usage_service.register_usage(
                provider_id=runtime.provider.id,
                model_id=runtime.model.id,
                execution_id=execution_id,
                input_tokens=provider_result.input_tokens,
                output_tokens=provider_result.output_tokens,
                estimated_cost=chunk_cost,
            )

            total_input_tokens += provider_result.input_tokens
            total_output_tokens += provider_result.output_tokens
            total_cost += chunk_cost
            providers_used.add(runtime.provider.slug)
            models_used.add(runtime.model.model_slug)
            chunk_stage_of_failure = "provider_response_validation"
            parsed_chunk = self.response_parser.parse(
                parser_strategy=parser_strategy,
                output_text=str(provider_result.output_text or "").strip(),
                output_schema=output_contract.output_schema,
            )
            chunk_stage_of_failure = "output_normalization"
            normalized_chunk = self.result_normalizer.normalize_text_chunk(parsed_chunk=parsed_chunk)
            output_chunks.append(normalized_chunk)
            if chunk_debug is not None:
                chunk_debug["parsed_output"] = parsed_chunk
                chunk_debug["normalized_output"] = normalized_chunk
                chunk_debug["stage_of_failure"] = ""
                debug_chunks.append(chunk_debug)
            self._log_execution_phase(
                phase="execution.pipeline.response_parse",
                message="Text chunk response parsed.",
                level="debug",
                execution_id=str(execution_id),
                chunk_index=chunk_index,
                parser_strategy=parser_strategy.value,
            )
            self._notify_progress_update(
                progress_callback=progress_callback,
                phase="processing_chunks",
                status_message="Processando conteudo textual no modelo.",
                processed_chunks=chunk_index,
                total_chunks=len(content_chunks),
                current_row=chunk_index,
            )

        self._notify_progress_update(
            progress_callback=progress_callback,
            phase="normalizing_output",
            status_message="Normalizando resposta final da execucao textual.",
            processed_chunks=len(content_chunks),
            total_chunks=len(content_chunks),
        )
        merged_output = self.result_normalizer.merge_text_chunks(chunks=output_chunks)
        self._notify_progress_update(
            progress_callback=progress_callback,
            phase="exporting_result",
            status_message="Gerando arquivo final da execucao textual.",
            processed_chunks=len(content_chunks),
            total_chunks=len(content_chunks),
        )
        formatted_output = self.result_formatter.format_text_output(
            execution_id=execution_id,
            output_text=merged_output,
            output_contract=output_contract,
            output_policy=self.output_policy,
        )
        self._log_execution_phase(
            phase="execution.pipeline.output_generation",
            message="Text output generated.",
            execution_id=str(execution_id),
            provider_calls=provider_calls,
            output_file_name=formatted_output.file_name,
            output_file_mime=formatted_output.mime_type,
            output_characters=len(merged_output),
        )
        auxiliary_files: list[GeneratedExecutionFile] = []
        if debug_enabled:
            auxiliary_files = self._build_text_debug_files(
                execution_id=execution_id,
                official_prompt=official_prompt,
                runtime=runtime,
                parser_strategy=parser_strategy,
                formatter_strategy=formatter_strategy,
                output_contract=output_contract,
                debug_chunks=debug_chunks,
                merged_output=merged_output,
            )
        return ProcessedOutput(
            content=formatted_output.content,
            file_name=formatted_output.file_name,
            mime_type=formatted_output.mime_type,
            total_input_tokens=total_input_tokens,
            total_output_tokens=total_output_tokens,
            total_cost=total_cost,
            providers_used=providers_used,
            models_used=models_used,
            provider_calls=provider_calls,
            processing_summary={
                "pipeline": "textual_single_pass",
                "chunk_count": len(content_chunks),
                "output_type": output_type.value,
                "parser_strategy": parser_strategy.value,
                "formatter_strategy": formatter_strategy.value,
                "output_contract_source": output_contract.source,
                "execution_profile": execution_profile.name,
            },
            auxiliary_files=auxiliary_files,
        )

    def _process_tabular_file(
        self,
        *,
        execution_id: UUID,
        file_path: str,
        file_name: str,
        official_prompt: str,
        runtime: ProviderRuntimeSelection,
        global_context: str | None = None,
        parser_strategy: ExecutionParserStrategy,
        formatter_strategy: ExecutionFormatterStrategy,
        output_type: ExecutionOutputType,
        output_contract: ExecutionOutputContract,
        execution_started_at: float,
        execution_profile: ExecutionOperationalProfile,
        debug_enabled: bool = False,
        automation_id: UUID | None = None,
        retry_count: int = 0,
        progress_callback: Callable[[ExecutionProgressUpdate], None] | None = None,
    ) -> ProcessedOutput:
        self._notify_progress_update(
            progress_callback=progress_callback,
            phase="reading_input",
            status_message="Lendo arquivo tabular de entrada.",
        )
        extension = Path(str(file_name or "")).suffix.lower()
        self._log_execution_phase(
            phase="execution.pipeline.file_read",
            message="Reading tabular input file.",
            execution_id=str(execution_id),
            file_extension=extension,
            parser_strategy=parser_strategy.value,
            output_type=output_type.value,
            has_global_context=bool(global_context),
        )
        file_bytes = self._read_input_file_bytes(file_path=file_path)
        input_rows, input_headers = self._load_tabular_rows(
            content=file_bytes,
            extension=extension,
        )
        if not input_rows:
            raise AppException(
                "Tabular file does not contain valid rows for processing.",
                status_code=422,
                code="tabular_file_without_rows",
            )

        total_rows = len(input_rows)
        self._enforce_execution_rows_profile_limit(
            execution_id=execution_id,
            total_rows=total_rows,
            execution_profile=execution_profile,
        )
        self._enforce_execution_rows_hard_limit(
            execution_id=execution_id,
            total_rows=total_rows,
        )
        provider_calls = 0
        self._log_execution_phase(
            phase="execution.pipeline.file_read.completed",
            message="Tabular input parsed and ready for row processing.",
            execution_id=str(execution_id),
            total_rows=total_rows,
            header_count=len(input_headers),
        )
        self._notify_progress_update(
            progress_callback=progress_callback,
            phase="processing_rows",
            status_message="Processamento linha a linha iniciado.",
            processed_rows=0,
            total_rows=total_rows,
        )
        self._log_execution_phase(
            phase="execution.pipeline.prompt_build",
            message="Tabular row-by-row prompt generation started.",
            execution_id=str(execution_id),
            total_rows=total_rows,
            has_global_context=bool(global_context),
        )

        total_input_tokens = 0
        total_output_tokens = 0
        total_cost = Decimal("0")
        providers_used: set[str] = set()
        models_used: set[str] = set()
        output_rows: list[dict[str, Any]] = []
        debug_rows: list[dict[str, Any]] = []
        successful_rows = 0
        failed_rows = 0
        output_schema = output_contract.output_schema
        prompt_strategy = self.tabular_prompt_strategy_resolver.resolve(output_schema=output_schema)
        detected_prompt_placeholders = prompt_strategy.detect_placeholders(str(official_prompt or ""))
        row_retry_count = max(int(retry_count or 0), 0)
        self._log_execution_phase(
            phase="execution.pipeline.prompt_template",
            message="Automation prompt template loaded for tabular hydration.",
            level="debug",
            execution_id=str(execution_id),
            prompt_template_preview=self._prompt_preview(official_prompt),
            detected_placeholders=list(detected_prompt_placeholders),
            prompt_field_columns=list(output_schema.prompt_field_columns.keys()),
        )
        tabular_context = self.result_normalizer.build_tabular_context(
            input_headers=input_headers,
            output_schema=output_schema,
        )

        for row in input_rows:
            row_index = int(row.get("row_index") or 0)
            row_values = row.get("values") if isinstance(row.get("values"), dict) else {}
            self._enforce_execution_time_profile_limit(
                execution_id=execution_id,
                execution_started_at=execution_started_at,
                phase="execution.pipeline.tabular_row_loop",
                execution_profile=execution_profile,
                row_index=row_index,
            )
            self._enforce_execution_time_hard_limit(
                execution_id=execution_id,
                execution_started_at=execution_started_at,
                phase="execution.pipeline.tabular_row_loop",
                row_index=row_index,
            )
            self._enforce_tabular_row_size_profile_limit(
                execution_id=execution_id,
                row_index=row_index,
                row_values=row_values,
                execution_profile=execution_profile,
            )
            self._enforce_tabular_row_size_hard_limit(
                execution_id=execution_id,
                row_index=row_index,
                row_values=row_values,
            )
            prompt_fields: dict[str, str] = {}
            prompt_field_sources: dict[str, str] = {}
            prompt_render = None
            prompt_input = ""
            row_stage_of_failure = "provider_call"
            debug_row: dict[str, Any] | None = None
            if debug_enabled:
                debug_row = {
                    "row_index": row_index,
                    "input_snapshot": row_values,
                    "canonical_fields": {},
                    "prompt_field_sources": {},
                    "placeholders_detected": [],
                    "placeholders_resolved": [],
                    "placeholders_unresolved": [],
                    "prompt_template": official_prompt,
                    "prompt_final": "",
                    **self._build_provider_debug_context(
                        runtime=runtime,
                        retry_count=row_retry_count,
                    ),
                    "provider_parameters": {
                        "temperature": getattr(settings, "temperature", None),
                        "max_tokens": getattr(settings, "max_tokens", None),
                    },
                    "response_raw_text": "",
                    "response_raw_payload": None,
                    "json_payload_cleaned": None,
                    "json_payload_parsed": None,
                    "parsed_output": None,
                    "normalized_output": None,
                    "projected_output_row": None,
                    "warnings": [],
                    "errors": [],
                }
            try:
                prompt_field_resolution = prompt_strategy.resolve_prompt_fields(row_values=row_values)
                prompt_fields = prompt_field_resolution.values
                prompt_field_sources = prompt_field_resolution.sources
                prompt_render = prompt_strategy.render_prompt_with_metadata(
                    official_prompt=official_prompt,
                    prompt_fields=prompt_fields,
                    global_context=global_context,
                    normalize_inline_text=self._normalize_inline_text,
                    assemble_prompt=self._assemble_prompt,
                    execution_profile=execution_profile,
                    field_sources=prompt_field_resolution.sources,
                )
                prompt_input = prompt_render.prompt_text
                if debug_row is not None:
                    debug_row["canonical_fields"] = prompt_fields
                    debug_row["prompt_field_sources"] = prompt_field_sources
                    debug_row["placeholders_detected"] = list(prompt_render.detected_placeholders)
                    debug_row["placeholders_resolved"] = list(prompt_render.resolved_placeholders)
                    debug_row["placeholders_unresolved"] = list(prompt_render.unresolved_placeholders)
                    debug_row["prompt_final"] = prompt_input
                self._log_execution_phase(
                    phase="execution.pipeline.prompt_hydration",
                    message="Tabular prompt hydrated with schema-resolved input fields.",
                    level="debug",
                    execution_id=str(execution_id),
                    row_index=row_index,
                    detected_placeholders=list(prompt_render.detected_placeholders),
                    resolved_placeholders=list(prompt_render.resolved_placeholders),
                    unresolved_placeholders=list(prompt_render.unresolved_placeholders),
                    prompt_field_sources=prompt_field_sources,
                    prompt_preview=self._prompt_preview(prompt_input),
                )
            except Exception as prompt_exc:
                if debug_row is not None:
                    debug_row["canonical_fields"] = prompt_fields
                    debug_row["prompt_field_sources"] = prompt_field_sources
                    debug_row["stage_of_failure"] = "prompt_hydration"
                    error_summary = self._summarize_execution_error(prompt_exc)
                    debug_row["provider_error_message"] = error_summary
                    debug_row["error_type"] = self._classify_execution_error_type(
                        exc=prompt_exc,
                        stage_of_failure="prompt_hydration",
                    )
                    debug_row["errors"] = [error_summary]
                    debug_rows.append(self._normalize_tabular_debug_row(debug_row=debug_row))
                raise

            output_row = self.result_normalizer.build_tabular_output_row(
                row_index=row_index,
                row_values=row_values,
                prompt_fields=prompt_fields,
                output_schema=output_schema,
                context=tabular_context,
            )

            try:
                provider_call_started_at = ""
                provider_call_finished_at = ""
                provider_call_duration_ms = 0
                self._enforce_provider_calls_profile_limit(
                    execution_id=execution_id,
                    provider_calls=provider_calls,
                    phase="execution.pipeline.provider_call",
                    execution_profile=execution_profile,
                    row_index=row_index,
                )
                self._enforce_provider_calls_hard_limit(
                    execution_id=execution_id,
                    provider_calls=provider_calls,
                    phase="execution.pipeline.provider_call",
                    row_index=row_index,
                )
                sanitized_prompt, _ = self._enforce_token_limit(
                    prompt=prompt_input,
                    provider_runtime=runtime,
                )
                request_payload_sanitized = self._build_provider_request_payload(
                    runtime=runtime,
                    prompt_input=sanitized_prompt,
                )
                token_limit_param_used = ""
                if isinstance(request_payload_sanitized, dict):
                    if "max_completion_tokens" in request_payload_sanitized:
                        token_limit_param_used = "max_completion_tokens"
                    elif "max_tokens" in request_payload_sanitized:
                        token_limit_param_used = "max_tokens"
                if debug_row is not None:
                    debug_row["prompt_final"] = sanitized_prompt
                    debug_row["request_payload_sanitized"] = request_payload_sanitized
                    debug_row["token_limit_param_used"] = token_limit_param_used
                    debug_row["api_family_resolved"] = "chat_completions"
                    if str(runtime.provider.slug or "").strip().lower() == "openai":
                        debug_row["request_profile_resolved"] = (
                            "gpt5_chat" if token_limit_param_used == "max_completion_tokens" else "legacy_chat"
                        )
                    debug_row["stage_of_failure"] = row_stage_of_failure
                row_client_request_id = f"{execution_id}:row:{row_index}:{uuid4().hex[:8]}"
                provider_call_started_at = datetime.now(timezone.utc).isoformat()
                provider_call_started_perf = perf_counter()
                if debug_row is not None:
                    debug_row["started_at"] = provider_call_started_at
                    debug_row["client_request_id"] = row_client_request_id
                self._log_execution_phase(
                    phase="execution.pipeline.provider_call",
                    message="Executing provider call for tabular row.",
                    level="debug",
                    execution_id=str(execution_id),
                    row_index=row_index,
                )
                provider_calls += 1
                provider_result = self._execute_with_runtime(
                    prompt_input=sanitized_prompt,
                    runtime=runtime,
                    client_request_id=row_client_request_id,
                )
                provider_call_finished_at = datetime.now(timezone.utc).isoformat()
                provider_call_duration_ms = max(int((perf_counter() - provider_call_started_perf) * 1000), 0)
                if debug_row is not None:
                    debug_row["finished_at"] = provider_call_finished_at
                    debug_row["duration_ms"] = provider_call_duration_ms
                    debug_row["stage_of_failure"] = "provider_response_validation"
                    debug_row["http_status_code"] = 200
                    debug_row["response_raw_text"] = str(provider_result.output_text or "")
                    debug_row["response_raw_payload"] = provider_result.raw_response
                line_cost = runtime.client.estimate_cost(
                    input_tokens=provider_result.input_tokens,
                    output_tokens=provider_result.output_tokens,
                    cost_input_per_1k_tokens=runtime.model.cost_input_per_1k_tokens,
                    cost_output_per_1k_tokens=runtime.model.cost_output_per_1k_tokens,
                )
                if total_cost + line_cost > Decimal(str(settings.max_cost_per_execution)):
                    raise AppException(
                        "Execution aborted due to estimated cost limit.",
                        status_code=422,
                        code="cost_limit_exceeded",
                        details={"max_cost_per_execution": settings.max_cost_per_execution},
                    )

                self.usage_service.register_usage(
                    provider_id=runtime.provider.id,
                    model_id=runtime.model.id,
                    execution_id=execution_id,
                    input_tokens=provider_result.input_tokens,
                    output_tokens=provider_result.output_tokens,
                    estimated_cost=line_cost,
                )

                row_stage_of_failure = "provider_response_validation"
                model_output_text = str(provider_result.output_text or "").strip()
                if not model_output_text:
                    raise AppException(
                        "Provider returned empty body.",
                        status_code=502,
                        code="provider_empty_output",
                        details={
                            "provider": str(runtime.provider.slug or "").strip().lower(),
                            "request_url": str(debug_row.get("request_url") if debug_row else ""),
                            "endpoint_name": str(debug_row.get("endpoint_name") if debug_row else ""),
                            "request_method": str(debug_row.get("request_method") if debug_row else ""),
                            "request_timeout_seconds": debug_row.get("request_timeout_seconds") if debug_row else None,
                            "request_payload_sanitized": request_payload_sanitized,
                            "started_at": provider_call_started_at,
                            "finished_at": provider_call_finished_at,
                            "duration_ms": provider_call_duration_ms,
                            "http_status_code": 200,
                            "status_code": 200,
                            "provider_error_message": "Provider returned empty body",
                            "provider_error_classification": "provider_empty_response",
                            "client_request_id": row_client_request_id,
                        },
                    )
                json_inspection: dict[str, Any] | None = None
                if parser_strategy == ExecutionParserStrategy.TABULAR_STRUCTURED:
                    schema_aliases = output_schema.structured_output_aliases if output_schema is not None else {}
                    json_inspection = self.response_parser.inspect_structured_output_json(
                        output_text=model_output_text,
                        structured_aliases=schema_aliases,
                    )

                row_stage_of_failure = "structured_parse"
                if debug_row is not None:
                    debug_row["stage_of_failure"] = row_stage_of_failure
                parsed_output = self.response_parser.parse(
                    parser_strategy=parser_strategy,
                    output_text=model_output_text,
                    output_schema=output_schema,
                )
                row_stage_of_failure = "output_normalization"
                if debug_row is not None:
                    debug_row["stage_of_failure"] = row_stage_of_failure
                normalized_output = self.result_normalizer.normalize_tabular_row_result(
                    parsed_output=parsed_output,
                    output_schema=output_schema,
                )
                row_stage_of_failure = "spreadsheet_projection"
                if debug_row is not None:
                    debug_row["stage_of_failure"] = row_stage_of_failure
                output_row.update(normalized_output)
                if debug_row is not None:
                    if json_inspection is not None:
                        debug_row["json_payload_cleaned"] = json_inspection.get("cleaned_payload")
                        debug_row["json_payload_parsed"] = json_inspection.get("parsed_json")
                        parse_error = str(json_inspection.get("parse_error") or "").strip()
                        if parse_error:
                            debug_row["warnings"] = [*list(debug_row.get("warnings") or []), parse_error]
                    debug_row["parsed_output"] = parsed_output
                    debug_row["normalized_output"] = normalized_output
                    debug_row["stage_of_failure"] = ""
                if output_schema.status_column:
                    output_row[output_schema.status_column] = "ok"
                if output_schema.error_column:
                    output_row[output_schema.error_column] = ""
                if debug_row is not None:
                    debug_row["projected_output_row"] = dict(output_row)
                successful_rows += 1
                self._log_execution_phase(
                    phase="execution.pipeline.response_parse",
                    message="Tabular row response parsed.",
                    level="debug",
                    execution_id=str(execution_id),
                    row_index=row_index,
                )

                total_input_tokens += provider_result.input_tokens
                total_output_tokens += provider_result.output_tokens
                total_cost += line_cost
                providers_used.add(runtime.provider.slug)
                models_used.add(runtime.model.model_slug)
            except Exception as row_exc:
                error_summary = self._summarize_execution_error(row_exc)
                error_type = self._classify_execution_error_type(
                    exc=row_exc,
                    stage_of_failure=row_stage_of_failure,
                )
                if debug_row is not None:
                    self._enrich_debug_with_error(
                        debug_item=debug_row,
                        exc=row_exc,
                        stage_of_failure=row_stage_of_failure,
                    )
                    debug_row["errors"] = [error_summary]
                    debug_row["projected_output_row"] = dict(output_row)
                if error_type.startswith("provider_") or row_stage_of_failure in {
                    "provider_call",
                    "provider_response_validation",
                }:
                    self._log_provider_error(
                        execution_id=execution_id,
                        row_index=row_index,
                        automation_id=automation_id,
                        runtime=runtime,
                        stage_of_failure=row_stage_of_failure,
                        exc=row_exc,
                    )
                if self._is_fatal_tabular_row_exception(row_exc):
                    if debug_row is not None:
                        debug_rows.append(self._normalize_tabular_debug_row(debug_row=debug_row))
                    self._log_execution_phase(
                        phase="execution.pipeline.row_by_row.fatal_error",
                        message="Fatal error while processing tabular row.",
                        level="error",
                        execution_id=str(execution_id),
                        row_index=row_index,
                        error_code=row_exc.payload.code if isinstance(row_exc, AppException) else "unexpected_error",
                        error_details=row_exc.payload.details if isinstance(row_exc, AppException) else None,
                    )
                    raise
                if output_schema.status_column:
                    output_row[output_schema.status_column] = "erro"
                if output_schema.error_column:
                    output_row[output_schema.error_column] = error_summary
                failed_rows += 1
                logger.warning(
                    "Failed to process tabular row; execution will continue for remaining rows.",
                    extra={
                        "execution_id": str(execution_id),
                        "row_index": row_index,
                        "automation_id": str(automation_id) if automation_id is not None else None,
                        "provider": str(runtime.provider.slug or "").strip().lower(),
                        "model": str(runtime.model.model_slug or "").strip(),
                        "status_code": debug_row.get("http_status_code") if debug_row is not None else None,
                        "error_type": error_type,
                        "request_id": debug_row.get("provider_request_id") if debug_row is not None else None,
                        "client_request_id": debug_row.get("client_request_id") if debug_row is not None else None,
                        "duration_ms": debug_row.get("duration_ms") if debug_row is not None else None,
                        "stage_of_failure": row_stage_of_failure,
                    },
                    exc_info=row_exc,
                )
                if debug_row is not None:
                    debug_row["projected_output_row"] = dict(output_row)
                    debug_rows.append(self._normalize_tabular_debug_row(debug_row=debug_row))
            else:
                if debug_row is not None:
                    debug_rows.append(self._normalize_tabular_debug_row(debug_row=debug_row))

            output_rows.append(output_row)
            self._notify_progress_update(
                progress_callback=progress_callback,
                phase="processing_rows",
                status_message="Processando linhas no modelo de IA.",
                processed_rows=len(output_rows),
                total_rows=total_rows,
                current_row=row_index,
            )

        self._notify_progress_update(
            progress_callback=progress_callback,
            phase="normalizing_output",
            status_message="Normalizando dados de saida da planilha.",
            processed_rows=len(output_rows),
            total_rows=total_rows,
        )
        output_columns = self.result_normalizer.build_tabular_output_columns(
            output_schema=output_schema,
            context=tabular_context,
        )
        self._notify_progress_update(
            progress_callback=progress_callback,
            phase="exporting_result",
            status_message="Gerando arquivo final com os resultados.",
            processed_rows=len(output_rows),
            total_rows=total_rows,
        )
        formatted_output = self.result_formatter.format_tabular_output(
            execution_id=execution_id,
            rows=output_rows,
            columns=output_columns,
            output_contract=output_contract,
            output_policy=self.output_policy,
            workbook_builder=self._build_tabular_workbook,
        )
        auxiliary_files: list[GeneratedExecutionFile] = []
        if debug_enabled:
            auxiliary_files = self._build_tabular_debug_files(
                execution_id=execution_id,
                runtime=runtime,
                parser_strategy=parser_strategy,
                formatter_strategy=formatter_strategy,
                output_contract=output_contract,
                debug_rows=debug_rows,
                retry_count=row_retry_count,
            )
        logger.info(
            "Tabular execution completed.",
            extra={
                "execution_id": str(execution_id),
                "processed_rows": len(output_rows),
                "successful_rows": successful_rows,
                "failed_rows": failed_rows,
                "provider_calls": provider_calls,
                "phase": "execution.pipeline.output_generation",
                "output_file_name": formatted_output.file_name,
                "output_file_mime": formatted_output.mime_type,
            },
        )
        return ProcessedOutput(
            content=formatted_output.content,
            file_name=formatted_output.file_name,
            mime_type=formatted_output.mime_type,
            total_input_tokens=total_input_tokens,
            total_output_tokens=total_output_tokens,
            total_cost=total_cost,
            providers_used=providers_used,
            models_used=models_used,
            provider_calls=provider_calls,
            processing_summary={
                "pipeline": "tabular_row_by_row_with_context" if global_context else "tabular_row_by_row",
                "total_rows": total_rows,
                "processed_rows": len(output_rows),
                "successful_rows": successful_rows,
                "failed_rows": failed_rows,
                "output_type": output_type.value,
                "parser_strategy": parser_strategy.value,
                "formatter_strategy": formatter_strategy.value,
                "output_contract_source": output_contract.source,
                "execution_profile": execution_profile.name,
            },
            auxiliary_files=auxiliary_files,
        )

    @staticmethod
    def _normalize_key(value: Any) -> str:
        raw = str(value or "").strip().lower()
        if not raw:
            return ""
        normalized = unicodedata.normalize("NFKD", raw)
        normalized = normalized.encode("ascii", "ignore").decode("ascii")
        normalized = re.sub(r"[^a-z0-9]+", "_", normalized).strip("_")
        return normalized

    @staticmethod
    def _normalize_tabular_value(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, datetime):
            return value.isoformat()
        return str(value).strip()

    def _load_tabular_rows(
        self,
        *,
        content: bytes,
        extension: str,
    ) -> tuple[list[dict[str, Any]], list[str]]:
        self._log_execution_phase(
            phase="execution.pipeline.file_parse",
            message="Parsing tabular input content.",
            file_extension=extension,
            input_bytes=len(content),
        )
        if extension == LEGACY_XLS_EXTENSION:
            raise AppException(
                "Legacy .xls files are not supported. Convert the spreadsheet to .xlsx.",
                status_code=422,
                code="xls_legacy_not_supported",
            )
        if extension == ".csv":
            rows, headers = self._load_csv_rows(content=content)
            self._log_execution_phase(
                phase="execution.pipeline.file_parse.completed",
                message="CSV content parsed.",
                file_extension=extension,
                total_rows=len(rows),
                header_count=len(headers),
            )
            return rows, headers
        if extension == ".xlsx":
            rows, headers = self._load_excel_rows(content=content)
            self._log_execution_phase(
                phase="execution.pipeline.file_parse.completed",
                message="Spreadsheet content parsed.",
                file_extension=extension,
                total_rows=len(rows),
                header_count=len(headers),
            )
            return rows, headers
        raise AppException(
            "Unsupported tabular file extension.",
            status_code=422,
            code="unsupported_tabular_extension",
        )

    def _load_csv_rows(self, *, content: bytes) -> tuple[list[dict[str, Any]], list[str]]:
        try:
            csv_text = content.decode("utf-8-sig", errors="ignore")
            reader = csv.DictReader(io.StringIO(csv_text))
        except Exception as exc:
            raise AppException(
                "Failed to parse CSV content.",
                status_code=422,
                code="tabular_file_parse_error",
            ) from exc

        raw_fieldnames = list(reader.fieldnames or [])
        if not raw_fieldnames:
            raise AppException(
                "CSV header row is empty.",
                status_code=422,
                code="tabular_file_header_missing",
            )
        duplicated_headers = self._collect_duplicate_headers(raw_fieldnames)
        if duplicated_headers:
            raise AppException(
                "CSV contains duplicated header names. Rename duplicated columns before upload.",
                status_code=422,
                code="tabular_file_duplicate_headers",
                details={"duplicate_headers": duplicated_headers},
            )
        headers = self._build_unique_headers(raw_fieldnames)

        rows: list[dict[str, Any]] = []
        for row_index, row in enumerate(reader, start=2):
            if not isinstance(row, dict):
                continue
            row_values: dict[str, str] = {}
            for position, header in enumerate(headers):
                source_key = raw_fieldnames[position]
                raw_value = row.get(source_key, "")
                if isinstance(raw_value, list):
                    raw_value = ", ".join(str(item) for item in raw_value if item is not None)
                row_values[header] = self._normalize_tabular_value(raw_value)
            if self._is_effectively_empty_row(row_values):
                continue
            rows.append({"row_index": row_index, "values": row_values})
        return rows, headers

    def _load_excel_rows(self, *, content: bytes) -> tuple[list[dict[str, Any]], list[str]]:
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(filename=io.BytesIO(content), read_only=True, data_only=True)
        except Exception as exc:
            raise AppException(
                "Failed to parse spreadsheet. Convert legacy .xls to .xlsx when necessary.",
                status_code=422,
                code="tabular_file_parse_error",
            ) from exc

        sheet = self._select_excel_sheet(workbook=workbook)
        header_values = self._extract_header_values(sheet)
        if not self._has_meaningful_headers(header_values):
            raise AppException(
                "Spreadsheet header row is empty.",
                status_code=422,
                code="tabular_file_header_missing",
            )
        headers = self._build_unique_headers(header_values)

        rows: list[dict[str, Any]] = []
        for row_index, values in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
            current_values = list(values or [])
            row_values: dict[str, str] = {}
            for index, header in enumerate(headers):
                cell_value = current_values[index] if index < len(current_values) else ""
                row_values[header] = self._normalize_tabular_value(cell_value)

            if self._is_effectively_empty_row(row_values):
                continue
            rows.append({"row_index": row_index, "values": row_values})
        return rows, headers

    def _select_excel_sheet(self, *, workbook: Any):  # type: ignore[no-untyped-def]
        worksheets = list(getattr(workbook, "worksheets", []) or [])
        if not worksheets:
            raise AppException(
                "Spreadsheet does not contain worksheets.",
                status_code=422,
                code="tabular_file_parse_error",
            )

        selected_sheet = worksheets[0]
        for sheet in worksheets:
            if self._has_meaningful_headers(self._extract_header_values(sheet)):
                selected_sheet = sheet
                break

        if selected_sheet is not worksheets[0]:
            logger.info(
                "Spreadsheet worksheet selected by non-empty header detection.",
                extra={"sheet_name": selected_sheet.title},
            )
        return selected_sheet

    @staticmethod
    def _extract_header_values(sheet) -> list[Any]:  # type: ignore[no-untyped-def]
        for header_row in sheet.iter_rows(min_row=1, max_row=1, values_only=True):
            return list(header_row or [])
        return []

    @classmethod
    def _has_meaningful_headers(cls, header_values: list[Any]) -> bool:
        return any(cls._normalize_tabular_value(value) for value in header_values)

    @classmethod
    def _build_unique_headers(cls, header_values: list[Any]) -> list[str]:
        if not header_values:
            return []

        headers: list[str] = []
        seen_counts: dict[str, int] = {}
        for index, header_value in enumerate(header_values):
            raw_header = cls._normalize_tabular_value(header_value) or f"coluna_{index + 1}"
            normalized_key = cls._normalize_key(raw_header) or f"coluna_{index + 1}"
            occurrences = seen_counts.get(normalized_key, 0) + 1
            seen_counts[normalized_key] = occurrences
            if occurrences > 1:
                headers.append(f"{raw_header}_{occurrences}")
            else:
                headers.append(raw_header)
        return headers

    @classmethod
    def _collect_duplicate_headers(cls, header_values: list[Any]) -> list[str]:
        occurrences: dict[str, int] = {}
        display_names: dict[str, str] = {}
        for raw_header in header_values:
            normalized = cls._normalize_tabular_value(raw_header)
            if not normalized:
                continue
            key = cls._normalize_key(normalized)
            if not key:
                continue
            occurrences[key] = occurrences.get(key, 0) + 1
            display_names[key] = normalized
        return sorted(display_names[key] for key, count in occurrences.items() if count > 1)

    @staticmethod
    def _is_effectively_empty_row(row_values: dict[str, Any]) -> bool:
        return not any(str(value or "").strip() for value in row_values.values())

    @staticmethod
    def _provider_endpoint_context(*, runtime: ProviderRuntimeSelection) -> dict[str, Any]:
        provider_slug = str(runtime.provider.slug or "").strip().lower()
        model_identifier = str(runtime.model.model_slug or "").strip()
        timeout_seconds = int(getattr(runtime.client, "timeout_seconds", settings.provider_timeout) or settings.provider_timeout)

        if provider_slug == "openai":
            base_url = str(getattr(runtime.client, "base_url", "https://api.openai.com/v1") or "https://api.openai.com/v1").rstrip("/")
            return {
                "endpoint_name": "chat_completions",
                "request_url": f"{base_url}/chat/completions",
                "request_method": "POST",
                "request_timeout_seconds": timeout_seconds,
            }
        if provider_slug == "anthropic":
            base_url = str(getattr(runtime.client, "base_url", "https://api.anthropic.com") or "https://api.anthropic.com").rstrip("/")
            return {
                "endpoint_name": "messages",
                "request_url": f"{base_url}/v1/messages",
                "request_method": "POST",
                "request_timeout_seconds": timeout_seconds,
            }
        if provider_slug == "gemini":
            config_json = getattr(runtime.client, "config_json", {}) or {}
            base_url = str(config_json.get("base_url") or "https://generativelanguage.googleapis.com").rstrip("/")
            api_version = str(config_json.get("api_version") or "v1beta").strip().strip("/") or "v1beta"
            model_id = model_identifier[7:] if model_identifier.startswith("models/") else model_identifier
            return {
                "endpoint_name": "generate_content",
                "request_url": f"{base_url}/{api_version}/models/{model_id}:generateContent",
                "request_method": "POST",
                "request_timeout_seconds": timeout_seconds,
            }
        return {
            "endpoint_name": "unknown",
            "request_url": "",
            "request_method": "POST",
            "request_timeout_seconds": timeout_seconds,
        }

    def _build_provider_request_payload(
        self,
        *,
        runtime: ProviderRuntimeSelection,
        prompt_input: str,
    ) -> dict[str, Any]:
        provider_slug = str(runtime.provider.slug or "").strip().lower()
        model_identifier = str(runtime.model.model_slug or "").strip()
        if provider_slug == "gemini":
            model_id = model_identifier[7:] if model_identifier.startswith("models/") else model_identifier
            payload: dict[str, Any] = {
                "model": model_identifier,
                "resolved_model_identifier": model_id,
                "input": prompt_input,
                "contents": [{"role": "user", "parts": [{"text": prompt_input}]}],
                "generationConfig": {
                    "temperature": getattr(settings, "temperature", None),
                    "maxOutputTokens": getattr(settings, "max_tokens", None),
                },
                "response_format": None,
                "tools": [],
                "metadata": {
                    "provider_slug": provider_slug,
                    "model_slug": model_identifier,
                },
            }
            return sanitize_provider_debug_payload(payload)

        payload = {
            "model": model_identifier,
            "resolved_model_identifier": model_identifier,
            "input": prompt_input,
            "messages": [{"role": "user", "content": prompt_input}],
            "temperature": getattr(settings, "temperature", None),
            "response_format": None,
            "tools": [],
            "metadata": {
                "provider_slug": provider_slug,
                "model_slug": model_identifier,
            },
        }
        token_limit_param = "max_tokens"
        if provider_slug == "openai":
            raw_model_metadata = getattr(runtime.model, "config_json", None)
            model_metadata = dict(raw_model_metadata) if isinstance(raw_model_metadata, dict) else {}
            for field_name in ("api_family", "token_limit_param", "request_profile", "supports_reasoning"):
                raw_value = getattr(runtime.model, field_name, None)
                if raw_value is not None and field_name not in model_metadata:
                    model_metadata[field_name] = raw_value

            explicit_token_param = str(model_metadata.get("token_limit_param") or "").strip().lower()
            explicit_request_profile = str(model_metadata.get("request_profile") or "").strip().lower()
            explicit_api_family = str(model_metadata.get("api_family") or "").strip().lower()
            supports_reasoning = str(model_metadata.get("supports_reasoning") or "").strip().lower()
            inferred_gpt5 = str(model_identifier or "").strip().lower().startswith("gpt-5")
            if explicit_token_param == "max_completion_tokens":
                token_limit_param = "max_completion_tokens"
            elif explicit_token_param == "max_tokens":
                token_limit_param = "max_tokens"
            elif explicit_request_profile in {"gpt5_chat", "gpt5_responses"}:
                token_limit_param = "max_completion_tokens"
            elif explicit_request_profile == "legacy_chat":
                token_limit_param = "max_tokens"
            elif explicit_api_family == "responses":
                token_limit_param = "max_completion_tokens"
            elif supports_reasoning in {"true", "1", "yes", "on"}:
                token_limit_param = "max_completion_tokens"
            elif inferred_gpt5:
                token_limit_param = "max_completion_tokens"

        if token_limit_param == "max_completion_tokens":
            payload["max_completion_tokens"] = getattr(settings, "max_tokens", None)
        else:
            payload["max_tokens"] = getattr(settings, "max_tokens", None)
        return sanitize_provider_debug_payload(payload)

    @staticmethod
    def _provider_debug_identity(*, runtime: ProviderRuntimeSelection) -> dict[str, Any]:
        provider_slug = str(runtime.provider.slug or "").strip().lower()
        provider_name = str(getattr(runtime.provider, "name", "") or provider_slug).strip() or provider_slug
        model_slug = str(runtime.model.model_slug or "").strip()
        model_name = str(getattr(runtime.model, "model_name", "") or model_slug).strip() or model_slug
        return {
            "provider_name": provider_name,
            "provider_slug": provider_slug,
            "model_name": model_name,
            "model_slug": model_slug,
            "resolved_model_identifier": model_slug or model_name,
        }

    def _build_provider_debug_context(
        self,
        *,
        runtime: ProviderRuntimeSelection,
        retry_count: int,
    ) -> dict[str, Any]:
        endpoint_context = self._provider_endpoint_context(runtime=runtime)
        request_timeout_raw = endpoint_context.get("request_timeout_seconds")
        try:
            request_timeout_seconds = int(request_timeout_raw) if request_timeout_raw is not None else None
        except (TypeError, ValueError):
            request_timeout_seconds = None
        return {
            **self._provider_debug_identity(runtime=runtime),
            "endpoint_name": endpoint_context.get("endpoint_name", ""),
            "request_url": endpoint_context.get("request_url", ""),
            "request_method": endpoint_context.get("request_method", ""),
            "request_timeout_seconds": request_timeout_seconds,
            "api_family_resolved": "",
            "request_profile_resolved": "",
            "token_limit_param_used": "",
            "client_request_id": "",
            "request_payload_sanitized": {},
            "started_at": "",
            "finished_at": "",
            "duration_ms": 0,
            "http_status_code": None,
            "provider_error_message": "",
            "provider_error_type": "",
            "provider_error_code": "",
            "provider_request_id": "",
            "provider_trace_id": "",
            "response_headers_relevantes": {},
            "response_body_text": "",
            "response_body_json": None,
            "transport_error_class": "",
            "transport_error_message": "",
            "retry_count": max(int(retry_count or 0), 0),
            "retried": bool((retry_count or 0) > 0),
            "stage_of_failure": "",
            "error_type": "",
        }

    @staticmethod
    def _provider_error_details(exc: Exception) -> dict[str, Any]:
        if isinstance(exc, AppException) and isinstance(exc.payload.details, dict):
            return dict(exc.payload.details)
        return {}

    def _classify_execution_error_type(
        self,
        *,
        exc: Exception,
        stage_of_failure: str,
    ) -> str:
        if stage_of_failure == "structured_parse":
            return "parse_error"
        if stage_of_failure == "output_normalization":
            return "parse_error"
        if stage_of_failure == "spreadsheet_projection":
            return "projection_error"

        if not isinstance(exc, AppException):
            return "unknown_provider_error"
        code = exc.payload.code
        details = self._provider_error_details(exc)
        if stage_of_failure == "provider_response_validation":
            if code == "provider_empty_output":
                return "provider_empty_response"
            if code == "provider_invalid_response":
                return "provider_non_json_error"
        if code == "provider_timeout":
            return "provider_timeout"
        if code == "provider_network_error":
            return "provider_connection_error"
        if code == "provider_empty_output":
            return "provider_empty_response"
        if code == "provider_invalid_response":
            return "provider_non_json_error"
        if code == "provider_http_error":
            provider_classification = str(details.get("provider_error_classification") or "").strip()
            if provider_classification:
                return provider_classification
            status_code = details.get("status_code") or details.get("http_status_code")
            try:
                status_code_int = int(status_code) if status_code is not None else None
            except (TypeError, ValueError):
                status_code_int = None
            if status_code_int in {401, 403}:
                return "provider_auth_error"
            if status_code_int == 429:
                return "provider_rate_limit"
            if status_code_int == 404:
                return "provider_unsupported_model"
            if status_code_int in {400, 422}:
                return "provider_invalid_request"
            return "provider_http_error"
        return "unknown_provider_error"

    def _summarize_execution_error(self, exc: Exception) -> str:
        if isinstance(exc, AppException):
            if exc.payload.code in {"provider_http_error", "provider_timeout", "provider_network_error"}:
                details = self._provider_error_details(exc)
                if details:
                    return summarize_provider_error_message(details=details)
            return exc.payload.message
        return str(exc)

    def _enrich_debug_with_error(
        self,
        *,
        debug_item: dict[str, Any],
        exc: Exception,
        stage_of_failure: str,
    ) -> None:
        details = self._provider_error_details(exc)
        debug_item["stage_of_failure"] = stage_of_failure
        debug_item["error_type"] = self._classify_execution_error_type(
            exc=exc,
            stage_of_failure=stage_of_failure,
        )
        debug_item["provider_error_message"] = str(
            details.get("provider_error_message")
            or self._summarize_execution_error(exc)
            or ""
        ).strip()
        debug_item["provider_error_type"] = str(details.get("provider_error_type") or "").strip()
        debug_item["provider_error_code"] = str(details.get("provider_error_code") or "").strip()
        debug_item["provider_request_id"] = str(details.get("provider_request_id") or "").strip()
        debug_item["provider_trace_id"] = str(details.get("provider_trace_id") or "").strip()
        debug_item["response_headers_relevantes"] = details.get("response_headers_relevantes") or {}
        debug_item["response_body_text"] = str(details.get("response_body_text") or "").strip()
        debug_item["response_body_json"] = details.get("response_body_json")
        debug_item["transport_error_class"] = str(details.get("transport_error_class") or "").strip()
        debug_item["transport_error_message"] = str(details.get("transport_error_message") or "").strip()
        debug_item["request_payload_sanitized"] = details.get("request_payload_sanitized") or debug_item.get(
            "request_payload_sanitized"
        ) or {}
        debug_item["api_family_resolved"] = str(
            details.get("api_family_resolved") or debug_item.get("api_family_resolved") or ""
        ).strip()
        debug_item["request_profile_resolved"] = str(
            details.get("request_profile_resolved") or debug_item.get("request_profile_resolved") or ""
        ).strip()
        debug_item["token_limit_param_used"] = str(
            details.get("token_limit_param_used") or debug_item.get("token_limit_param_used") or ""
        ).strip()
        debug_item["client_request_id"] = str(details.get("client_request_id") or debug_item.get("client_request_id") or "").strip()
        debug_item["started_at"] = str(details.get("started_at") or debug_item.get("started_at") or "").strip()
        debug_item["finished_at"] = str(details.get("finished_at") or debug_item.get("finished_at") or "").strip()
        debug_item["request_url"] = str(details.get("request_url") or debug_item.get("request_url") or "").strip()
        debug_item["endpoint_name"] = str(details.get("endpoint_name") or debug_item.get("endpoint_name") or "").strip()
        debug_item["request_method"] = str(details.get("request_method") or debug_item.get("request_method") or "").strip()
        timeout_value = details.get("request_timeout_seconds", debug_item.get("request_timeout_seconds"))
        try:
            debug_item["request_timeout_seconds"] = int(timeout_value) if timeout_value is not None else None
        except (TypeError, ValueError):
            debug_item["request_timeout_seconds"] = None
        status_value = details.get("status_code", details.get("http_status_code"))
        try:
            debug_item["http_status_code"] = int(status_value) if status_value is not None else None
        except (TypeError, ValueError):
            debug_item["http_status_code"] = None
        duration_value = details.get("duration_ms")
        try:
            debug_item["duration_ms"] = max(int(duration_value), 0) if duration_value is not None else debug_item.get("duration_ms")
        except (TypeError, ValueError):
            pass

    def _log_provider_error(
        self,
        *,
        execution_id: UUID,
        row_index: int | None,
        automation_id: UUID | None,
        runtime: ProviderRuntimeSelection,
        stage_of_failure: str,
        exc: Exception,
    ) -> None:
        details = self._provider_error_details(exc)
        error_type = self._classify_execution_error_type(
            exc=exc,
            stage_of_failure=stage_of_failure,
        )
        status_code = details.get("status_code") or details.get("http_status_code")
        duration_ms = details.get("duration_ms")
        request_id = details.get("provider_request_id")
        client_request_id = details.get("client_request_id")
        logger.warning(
            "Provider call failed during execution.",
            extra={
                "event": "provider_call_failed",
                "execution_id": str(execution_id),
                "row_index": int(row_index) if row_index is not None else None,
                "automation_id": str(automation_id) if automation_id is not None else None,
                "provider": str(runtime.provider.slug or "").strip().lower(),
                "model": str(runtime.model.model_slug or "").strip(),
                "status_code": int(status_code) if isinstance(status_code, int) or str(status_code).isdigit() else None,
                "error_type": error_type,
                "request_id": str(request_id or "").strip() or None,
                "client_request_id": str(client_request_id or "").strip() or None,
                "duration_ms": int(duration_ms) if isinstance(duration_ms, int) or str(duration_ms).isdigit() else None,
                "stage_of_failure": stage_of_failure,
            },
        )

    @classmethod
    def _normalize_tabular_debug_row(cls, *, debug_row: dict[str, Any]) -> dict[str, Any]:
        return {
            "row_index": int(debug_row.get("row_index") or 0),
            "input_snapshot": cls._debug_json_cell(debug_row.get("input_snapshot")),
            "canonical_fields": cls._debug_json_cell(debug_row.get("canonical_fields")),
            "prompt_field_sources": cls._debug_json_cell(debug_row.get("prompt_field_sources")),
            "placeholders_detected": cls._debug_json_cell(debug_row.get("placeholders_detected")),
            "placeholders_resolved": cls._debug_json_cell(debug_row.get("placeholders_resolved")),
            "placeholders_unresolved": cls._debug_json_cell(debug_row.get("placeholders_unresolved")),
            "prompt_template": cls._debug_text_cell(debug_row.get("prompt_template")),
            "prompt_final": cls._debug_text_cell(debug_row.get("prompt_final")),
            "provider_name": cls._debug_text_cell(debug_row.get("provider_name")),
            "provider_slug": cls._debug_text_cell(debug_row.get("provider_slug")),
            "model_name": cls._debug_text_cell(debug_row.get("model_name")),
            "model_slug": cls._debug_text_cell(debug_row.get("model_slug")),
            "resolved_model_identifier": cls._debug_text_cell(debug_row.get("resolved_model_identifier")),
            "request_url": cls._debug_text_cell(debug_row.get("request_url")),
            "endpoint_name": cls._debug_text_cell(debug_row.get("endpoint_name")),
            "request_method": cls._debug_text_cell(debug_row.get("request_method")),
            "request_timeout_seconds": cls._debug_text_cell(debug_row.get("request_timeout_seconds")),
            "api_family_resolved": cls._debug_text_cell(debug_row.get("api_family_resolved")),
            "request_profile_resolved": cls._debug_text_cell(debug_row.get("request_profile_resolved")),
            "token_limit_param_used": cls._debug_text_cell(debug_row.get("token_limit_param_used")),
            "client_request_id": cls._debug_text_cell(debug_row.get("client_request_id")),
            "request_payload_sanitized": cls._debug_json_cell(debug_row.get("request_payload_sanitized")),
            "started_at": cls._debug_text_cell(debug_row.get("started_at")),
            "finished_at": cls._debug_text_cell(debug_row.get("finished_at")),
            "duration_ms": cls._debug_text_cell(debug_row.get("duration_ms")),
            "http_status_code": cls._debug_text_cell(debug_row.get("http_status_code")),
            "provider_error_message": cls._debug_text_cell(debug_row.get("provider_error_message")),
            "provider_error_type": cls._debug_text_cell(debug_row.get("provider_error_type")),
            "provider_error_code": cls._debug_text_cell(debug_row.get("provider_error_code")),
            "provider_request_id": cls._debug_text_cell(debug_row.get("provider_request_id")),
            "provider_trace_id": cls._debug_text_cell(debug_row.get("provider_trace_id")),
            "response_headers_relevantes": cls._debug_json_cell(debug_row.get("response_headers_relevantes")),
            "response_body_text": cls._debug_text_cell(debug_row.get("response_body_text")),
            "response_body_json": cls._debug_json_cell(debug_row.get("response_body_json")),
            "transport_error_class": cls._debug_text_cell(debug_row.get("transport_error_class")),
            "transport_error_message": cls._debug_text_cell(debug_row.get("transport_error_message")),
            "retry_count": cls._debug_text_cell(debug_row.get("retry_count")),
            "retried": cls._debug_text_cell(debug_row.get("retried")),
            "stage_of_failure": cls._debug_text_cell(debug_row.get("stage_of_failure")),
            "error_type": cls._debug_text_cell(debug_row.get("error_type")),
            "provider_parameters": cls._debug_json_cell(debug_row.get("provider_parameters")),
            "response_raw_text": cls._debug_text_cell(debug_row.get("response_raw_text")),
            "response_raw_payload": cls._debug_json_cell(debug_row.get("response_raw_payload")),
            "json_payload_cleaned": cls._debug_text_cell(debug_row.get("json_payload_cleaned")),
            "json_payload_parsed": cls._debug_json_cell(debug_row.get("json_payload_parsed")),
            "parsed_output": cls._debug_json_cell(debug_row.get("parsed_output")),
            "normalized_output": cls._debug_json_cell(debug_row.get("normalized_output")),
            "projected_output_row": cls._debug_json_cell(debug_row.get("projected_output_row")),
            "warnings": cls._debug_json_cell(debug_row.get("warnings")),
            "errors": cls._debug_json_cell(debug_row.get("errors")),
        }

    def _build_tabular_debug_files(
        self,
        *,
        execution_id: UUID,
        runtime: ProviderRuntimeSelection,
        parser_strategy: ExecutionParserStrategy,
        formatter_strategy: ExecutionFormatterStrategy,
        output_contract: ExecutionOutputContract,
        debug_rows: list[dict[str, Any]],
        retry_count: int = 0,
    ) -> list[GeneratedExecutionFile]:
        columns = [
            "row_index",
            "input_snapshot",
            "canonical_fields",
            "prompt_field_sources",
            "placeholders_detected",
            "placeholders_resolved",
            "placeholders_unresolved",
            "prompt_template",
            "prompt_final",
            "provider_name",
            "provider_slug",
            "model_name",
            "model_slug",
            "resolved_model_identifier",
            "request_url",
            "endpoint_name",
            "request_method",
            "request_timeout_seconds",
            "api_family_resolved",
            "request_profile_resolved",
            "token_limit_param_used",
            "client_request_id",
            "request_payload_sanitized",
            "started_at",
            "finished_at",
            "duration_ms",
            "http_status_code",
            "provider_error_message",
            "provider_error_type",
            "provider_error_code",
            "provider_request_id",
            "provider_trace_id",
            "response_headers_relevantes",
            "response_body_text",
            "response_body_json",
            "transport_error_class",
            "transport_error_message",
            "retry_count",
            "retried",
            "stage_of_failure",
            "error_type",
            "provider_parameters",
            "response_raw_text",
            "response_raw_payload",
            "json_payload_cleaned",
            "json_payload_parsed",
            "parsed_output",
            "normalized_output",
            "projected_output_row",
            "warnings",
            "errors",
        ]
        endpoint_context = self._provider_endpoint_context(runtime=runtime)
        meta_row = self._normalize_tabular_debug_row(
            debug_row={
                "row_index": 0,
                "input_snapshot": {},
                "canonical_fields": {},
                "prompt_field_sources": {},
                "placeholders_detected": [],
                "placeholders_resolved": [],
                "placeholders_unresolved": [],
                "prompt_template": "",
                "prompt_final": "",
                "provider_name": str(getattr(runtime.provider, "name", "") or runtime.provider.slug),
                "provider_slug": runtime.provider.slug,
                "model_name": str(getattr(runtime.model, "model_name", "") or runtime.model.model_slug),
                "model_slug": runtime.model.model_slug,
                "resolved_model_identifier": runtime.model.model_slug,
                "request_url": endpoint_context.get("request_url"),
                "endpoint_name": endpoint_context.get("endpoint_name"),
                "request_method": endpoint_context.get("request_method"),
                "request_timeout_seconds": endpoint_context.get(
                    "request_timeout_seconds",
                    getattr(runtime.client, "timeout_seconds", settings.provider_timeout),
                ),
                "api_family_resolved": "",
                "request_profile_resolved": "",
                "token_limit_param_used": "",
                "client_request_id": "",
                "request_payload_sanitized": {},
                "started_at": "",
                "finished_at": "",
                "duration_ms": 0,
                "http_status_code": None,
                "provider_error_message": "",
                "provider_error_type": "",
                "provider_error_code": "",
                "provider_request_id": "",
                "provider_trace_id": "",
                "response_headers_relevantes": {},
                "response_body_text": "",
                "response_body_json": None,
                "transport_error_class": "",
                "transport_error_message": "",
                "retry_count": max(int(retry_count or 0), 0),
                "retried": bool((retry_count or 0) > 0),
                "stage_of_failure": "",
                "error_type": "",
                "provider_parameters": {
                    "temperature": getattr(settings, "temperature", None),
                    "max_tokens": getattr(settings, "max_tokens", None),
                    "parser_strategy": parser_strategy.value,
                    "formatter_strategy": formatter_strategy.value,
                    "output_contract_source": output_contract.source,
                },
                "response_raw_text": "",
                "response_raw_payload": None,
                "json_payload_cleaned": None,
                "json_payload_parsed": None,
                "parsed_output": None,
                "normalized_output": None,
                "projected_output_row": None,
                "warnings": [],
                "errors": [],
            }
        )
        rows = [meta_row, *debug_rows]
        content = self._build_tabular_workbook(
            rows=rows,
            columns=columns,
            worksheet_name="debug_execucao",
        )
        return [
            GeneratedExecutionFile(
                file_type="debug",
                file_name=f"debug_{execution_id}.xlsx",
                mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                content=content,
            )
        ]

    def _build_text_debug_files(
        self,
        *,
        execution_id: UUID,
        official_prompt: str,
        runtime: ProviderRuntimeSelection,
        parser_strategy: ExecutionParserStrategy,
        formatter_strategy: ExecutionFormatterStrategy,
        output_contract: ExecutionOutputContract,
        debug_chunks: list[dict[str, Any]],
        merged_output: str,
    ) -> list[GeneratedExecutionFile]:
        payload = {
            "execution_id": str(execution_id),
            "pipeline": "textual_single_pass",
            "provider_name": str(getattr(runtime.provider, "name", "") or runtime.provider.slug),
            "provider_slug": runtime.provider.slug,
            "model_name": str(getattr(runtime.model, "model_name", "") or runtime.model.model_slug),
            "model_slug": runtime.model.model_slug,
            "resolved_model_identifier": runtime.model.model_slug,
            "provider_parameters": {
                "temperature": getattr(settings, "temperature", None),
                "max_tokens": getattr(settings, "max_tokens", None),
            },
            "parser_strategy": parser_strategy.value,
            "formatter_strategy": formatter_strategy.value,
            "output_contract_source": output_contract.source,
            "prompt_template": self._debug_text_cell(official_prompt, max_characters=12000),
            "chunks": debug_chunks,
            "merged_output_preview": self._debug_text_cell(merged_output, max_characters=12000),
        }
        content = json.dumps(payload, ensure_ascii=False, default=str, indent=2).encode("utf-8")
        return [
            GeneratedExecutionFile(
                file_type="debug",
                file_name=f"debug_{execution_id}.json",
                mime_type="application/json",
                content=content,
            )
        ]

    @staticmethod
    def _debug_text_cell(value: Any, *, max_characters: int = 30000) -> str:
        if value is None:
            return ""
        normalized = str(value).strip()
        if not normalized:
            return ""
        if len(normalized) <= max_characters:
            return normalized
        return f"{normalized[:max_characters].rstrip()}...[truncated]"

    @classmethod
    def _debug_json_cell(cls, value: Any, *, max_characters: int = 30000) -> str:
        try:
            serialized = json.dumps(value, ensure_ascii=False, default=str)
        except Exception:
            serialized = str(value or "")
        return cls._debug_text_cell(serialized, max_characters=max_characters)

    @staticmethod
    def _build_tabular_workbook(
        *,
        rows: list[dict[str, Any]],
        columns: list[str],
        worksheet_name: str = "resultado",
    ) -> bytes:
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        normalized_sheet_name = str(worksheet_name or "resultado").strip()
        sheet.title = normalized_sheet_name[:31] or "resultado"
        sheet.append(columns)
        for row in rows:
            sheet.append([row.get(column, "") for column in columns])
        sheet.freeze_panes = "A2"

        output = io.BytesIO()
        workbook.save(output)
        return output.getvalue()

    def _read_input_file_bytes(self, *, file_path: str) -> bytes:
        with self.file_service.storage.open_file(file_path) as handle:
            return handle.read()

    def _read_input_file_content(self, *, file_path: str, file_name: str) -> str:
        extension = Path(file_name).suffix.lower()
        with self.file_service.storage.open_file(file_path) as handle:
            raw_bytes = handle.read()
            if extension == ".pdf":
                return self._extract_pdf_text(raw_bytes)
            if extension == ".xlsx":
                return self._extract_xlsx_text(raw_bytes)
            if extension == ".docx":
                return self._extract_docx_text(raw_bytes)
            if extension == ".doc":
                return self._extract_legacy_doc_text(raw_bytes)
            return raw_bytes.decode("utf-8", errors="ignore")

    @staticmethod
    def _extract_pdf_text(content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except Exception:
            return content[:8000].decode("utf-8", errors="ignore")

        try:
            import io

            reader = PdfReader(io.BytesIO(content))
            pages = []
            for page in reader.pages[:20]:
                pages.append(page.extract_text() or "")
            return "\n".join(pages).strip()
        except Exception:
            return content[:8000].decode("utf-8", errors="ignore")

    @staticmethod
    def _extract_xlsx_text(content: bytes) -> str:
        try:
            from openpyxl import load_workbook
        except Exception:
            return content[:8000].decode("utf-8", errors="ignore")

    @staticmethod
    def _extract_docx_text(content: bytes) -> str:
        try:
            from docx import Document
        except Exception:
            return content[:8000].decode("utf-8", errors="ignore")

        try:
            document = Document(io.BytesIO(content))
            blocks: list[str] = []
            for paragraph in document.paragraphs:
                text = paragraph.text.strip()
                if text:
                    blocks.append(text)
            for table in document.tables:
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if values:
                        blocks.append(" | ".join(values))
            return "\n".join(blocks).strip()
        except Exception:
            return content[:8000].decode("utf-8", errors="ignore")

    @staticmethod
    def _extract_legacy_doc_text(content: bytes) -> str:
        candidates: list[str] = []
        for encoding in ("utf-16-le", "cp1252", "latin-1"):
            try:
                decoded = content.decode(encoding, errors="ignore")
            except Exception:
                continue
            chunks = re.findall(r"[^\x00-\x1f\x7f-\x9f]{4,}", decoded)
            candidates.extend(chunk.strip() for chunk in chunks if chunk.strip())

        seen: set[str] = set()
        lines: list[str] = []
        for candidate in candidates:
            normalized = re.sub(r"\s+", " ", candidate).strip()
            if len(normalized) < 4 or normalized in seen:
                continue
            seen.add(normalized)
            lines.append(normalized)
            if len(lines) >= 1000:
                break
        return "\n".join(lines).strip()

        try:
            import io

            workbook = load_workbook(filename=io.BytesIO(content), read_only=True, data_only=True)
            rows_text: list[str] = []
            for sheet in workbook.worksheets[:3]:
                for row in sheet.iter_rows(min_row=1, max_row=500, values_only=True):
                    values = [str(value) for value in row if value is not None]
                    if values:
                        rows_text.append(", ".join(values))
            return "\n".join(rows_text).strip()
        except Exception:
            return content[:8000].decode("utf-8", errors="ignore")

    @staticmethod
    def _normalize_inline_text(value: str) -> str:
        return re.sub(r"\s+", " ", str(value or "")).strip()

    def _normalize_prompt_block_text(self, value: str, *, dedupe_consecutive_lines: bool) -> str:
        raw_text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
        normalized_lines: list[str] = []
        previous_key = ""
        for raw_line in raw_text.split("\n"):
            line = self._normalize_inline_text(raw_line)
            if not line:
                continue
            line_key = line.casefold()
            if dedupe_consecutive_lines and line_key == previous_key:
                continue
            normalized_lines.append(line)
            previous_key = line_key
        return "\n".join(normalized_lines).strip()

    @staticmethod
    def _truncate_text_at_boundary(
        *,
        text: str,
        max_characters: int,
        truncation_notice: str | None = None,
    ) -> tuple[str, bool]:
        normalized = str(text or "").strip()
        if max_characters <= 0:
            return "", bool(normalized)
        if len(normalized) <= max_characters:
            return normalized, False

        suffix = f"\n\n[{truncation_notice}]" if truncation_notice else ""
        available = max_characters - len(suffix)
        if available <= 0:
            fallback = suffix.strip() if suffix else normalized[:max_characters].rstrip()
            return fallback, True

        candidate = normalized[:available]
        breakpoints = [
            candidate.rfind("\n"),
            candidate.rfind(". "),
            candidate.rfind("; "),
            candidate.rfind(", "),
            candidate.rfind(" "),
        ]
        breakpoint = max(breakpoints)
        if breakpoint >= int(available * 0.6):
            candidate = candidate[:breakpoint]
        candidate = candidate.rstrip()
        if not candidate:
            candidate = normalized[:available].rstrip()
        merged = f"{candidate}{suffix}" if suffix else candidate
        return merged.strip(), True

    @staticmethod
    def _context_fingerprint(value: str) -> str:
        normalized = re.sub(r"\s+", " ", str(value or "")).strip().casefold()
        if len(normalized) > 5000:
            return normalized[:5000]
        return normalized

    @staticmethod
    def _context_type_priority(*, file_name: str) -> int:
        extension = Path(str(file_name or "")).suffix.lower()
        if extension in CONTEXT_STRUCTURED_EXTENSIONS:
            return 0
        if extension in CONTEXT_RAW_EXTENSIONS:
            return 1
        return 2

    @staticmethod
    def _describe_context_type(*, file_name: str) -> str:
        extension = Path(str(file_name or "")).suffix.lower()
        if extension in CONTEXT_STRUCTURED_EXTENSIONS:
            return "estruturado"
        if extension in CONTEXT_RAW_EXTENSIONS:
            return "bruto"
        return "outro"

    def _prioritize_context_inputs(self, *, context_inputs: list[EngineExecutionInput]) -> list[EngineExecutionInput]:
        return sorted(
            context_inputs,
            key=lambda item: (
                int(item.order_index),
                self._context_type_priority(file_name=item.file_name),
                str(item.file_name or "").lower(),
            ),
        )

    def _context_max_characters_per_file(self, *, execution_profile: ExecutionOperationalProfile) -> int:
        profile_limit = self._safe_hard_limit(execution_profile.max_context_file_characters, fallback=1)
        global_limit = self._safe_hard_limit(getattr(settings, "max_context_file_characters", 0), fallback=1)
        return min(profile_limit, global_limit)

    def _context_max_characters_total(self, *, execution_profile: ExecutionOperationalProfile) -> int:
        profile_limit = self._safe_hard_limit(execution_profile.max_context_characters, fallback=1)
        global_limit = self._safe_hard_limit(getattr(settings, "max_context_characters", 0), fallback=1)
        return min(profile_limit, global_limit)

    def _prompt_max_characters(self, *, execution_profile: ExecutionOperationalProfile) -> int:
        profile_limit = self._safe_hard_limit(execution_profile.max_prompt_characters, fallback=1)
        global_limit = self._safe_hard_limit(getattr(settings, "max_prompt_characters", 0), fallback=1)
        return min(profile_limit, global_limit)

    def _assemble_prompt(
        self,
        *,
        instruction_text: str,
        row_data: dict[str, str] | None = None,
        auxiliary_context: str | None = None,
        execution_profile: ExecutionOperationalProfile,
    ) -> str:
        normalized_instruction = self._normalize_prompt_block_text(
            instruction_text,
            dedupe_consecutive_lines=False,
        )
        sections = [
            f"{PROMPT_SECTION_INSTRUCTION}\n{normalized_instruction or '(instrucao principal ausente)'}"
        ]

        if row_data:
            row_lines: list[str] = []
            for field, value in row_data.items():
                normalized_value = self._normalize_inline_text(value)
                row_lines.append(f"{field}: {normalized_value}")
            sections.append(f"{PROMPT_SECTION_ROW_DATA}\n" + "\n".join(row_lines))

        normalized_context = self._normalize_prompt_block_text(
            auxiliary_context or "",
            dedupe_consecutive_lines=True,
        )
        if normalized_context:
            sections.append(f"{PROMPT_SECTION_CONTEXT}\n{normalized_context}")

        prompt_text = "\n\n".join(section for section in sections if section).strip()
        max_prompt_characters = self._prompt_max_characters(execution_profile=execution_profile)
        truncated_prompt, was_truncated = self._truncate_text_at_boundary(
            text=prompt_text,
            max_characters=max_prompt_characters,
            truncation_notice=f"prompt truncado para {max_prompt_characters} caracteres",
        )
        if was_truncated:
            logger.warning(
                "Prompt assembly exceeded configured size limit and was truncated.",
                extra={
                    "event": "prompt_truncated",
                    "max_prompt_characters": max_prompt_characters,
                },
            )
        return truncated_prompt

    @staticmethod
    def _prompt_preview(prompt_text: str, *, max_characters: int = 280) -> str:
        normalized = re.sub(r"\s+", " ", str(prompt_text or "")).strip()
        if len(normalized) <= max_characters:
            return normalized
        return f"{normalized[:max_characters].rstrip()}..."

    def _build_global_context_text(
        self,
        *,
        context_inputs: list[EngineExecutionInput],
        execution_profile: ExecutionOperationalProfile,
    ) -> str | None:
        if not context_inputs:
            return None

        self._log_execution_phase(
            phase="execution.pipeline.context_build",
            message="Building global textual context for tabular execution.",
            context_file_count=len(context_inputs),
        )
        prioritized_contexts = self._prioritize_context_inputs(context_inputs=context_inputs)
        max_context_chars_total = self._context_max_characters_total(execution_profile=execution_profile)
        max_context_chars_per_file = self._context_max_characters_per_file(execution_profile=execution_profile)

        blocks: list[str] = []
        dedupe_index: set[str] = set()
        duplicate_contexts = 0
        truncated_context_files = 0

        for context_input in prioritized_contexts:
            file_name = str(context_input.file_name or "arquivo_sem_nome")
            content = self._read_input_file_content(
                file_path=context_input.file_path,
                file_name=file_name,
            )
            normalized_content = self._normalize_prompt_block_text(
                content,
                dedupe_consecutive_lines=True,
            )
            safe_content = normalized_content or "(arquivo sem conteudo textual)"
            safe_content, was_file_truncated = self._truncate_text_at_boundary(
                text=safe_content,
                max_characters=max_context_chars_per_file,
                truncation_notice=f"contexto truncado para {max_context_chars_per_file} caracteres",
            )
            if was_file_truncated:
                truncated_context_files += 1

            fingerprint = self._context_fingerprint(safe_content)
            if fingerprint in dedupe_index:
                duplicate_contexts += 1
                continue
            dedupe_index.add(fingerprint)

            context_position = len(blocks) + 1
            context_type = self._describe_context_type(file_name=file_name)
            blocks.append(
                f"[Contexto {context_position} - {file_name} | ordem={int(context_input.order_index)} | tipo={context_type}]\n"
                f"{safe_content}"
            )

        merged_context = "\n\n".join(blocks).strip()
        if not merged_context:
            return None

        merged_context, was_context_truncated = self._truncate_text_at_boundary(
            text=merged_context,
            max_characters=max_context_chars_total,
            truncation_notice=f"contexto truncado para {max_context_chars_total} caracteres",
        )
        if was_context_truncated:
            logger.warning(
                "Global context exceeded configured limit and was truncated.",
                extra={
                    "event": "context_truncated",
                    "context_files": len(blocks),
                    "max_context_chars": max_context_chars_total,
                },
            )

        self._log_execution_phase(
            phase="execution.pipeline.context_build.completed",
            message="Global textual context prepared.",
            context_file_count=len(blocks),
            context_characters=len(merged_context),
            duplicated_context_files=duplicate_contexts,
            truncated_context_files=truncated_context_files,
            max_context_chars=max_context_chars_total,
        )
        return merged_context

    def _combine_textual_inputs(
        self,
        *,
        execution_inputs: list[EngineExecutionInput],
        execution_profile: ExecutionOperationalProfile,
    ) -> str:
        self._log_execution_phase(
            phase="execution.pipeline.multi_text_combine",
            message="Combining multiple textual inputs into a single analysis payload.",
            input_file_count=len(execution_inputs),
        )
        blocks: list[str] = []
        dedupe_index: set[str] = set()
        max_context_chars_per_file = self._context_max_characters_per_file(execution_profile=execution_profile)

        for execution_input in execution_inputs:
            file_name = str(execution_input.file_name or "arquivo_sem_nome")
            content = self._read_input_file_content(
                file_path=execution_input.file_path,
                file_name=file_name,
            )
            normalized_content = self._normalize_prompt_block_text(
                content,
                dedupe_consecutive_lines=True,
            )
            safe_content = normalized_content or "(arquivo sem conteudo textual)"
            if execution_input.role == INPUT_ROLE_CONTEXT:
                safe_content, _ = self._truncate_text_at_boundary(
                    text=safe_content,
                    max_characters=max_context_chars_per_file,
                    truncation_notice=f"contexto truncado para {max_context_chars_per_file} caracteres",
                )

            fingerprint = self._context_fingerprint(safe_content)
            if fingerprint in dedupe_index:
                continue
            dedupe_index.add(fingerprint)
            document_position = len(blocks) + 1
            blocks.append(
                f"[Documento {document_position} - {file_name}]\n{safe_content}"
            )
        merged_text = "\n\n".join(blocks).strip()
        self._log_execution_phase(
            phase="execution.pipeline.multi_text_combine.completed",
            message="Multiple textual inputs combined.",
            input_file_count=len(blocks),
            combined_characters=len(merged_text),
        )
        return merged_text

    def _chunk_content(self, content: str) -> list[str]:
        normalized = self._normalize_prompt_block_text(content, dedupe_consecutive_lines=False)
        if not normalized:
            return ["(arquivo sem conteudo textual)"]

        normalized, was_truncated = self._truncate_text_at_boundary(
            text=normalized,
            max_characters=settings.max_input_characters,
            truncation_notice=None,
        )
        if was_truncated:
            logger.warning(
                "Input content exceeded max_input_characters; truncating before chunking.",
                extra={"event": "content_truncated"},
            )

        if len(normalized) <= settings.chunk_size_characters:
            return [normalized]

        chunks: list[str] = []
        start = 0
        while start < len(normalized):
            raw_end = min(start + settings.chunk_size_characters, len(normalized))
            end = raw_end
            if raw_end < len(normalized):
                preferred_break = max(
                    normalized.rfind("\n", start, raw_end),
                    normalized.rfind(" ", start, raw_end),
                )
                if preferred_break > start + int(settings.chunk_size_characters * 0.6):
                    end = preferred_break
            if end <= start:
                end = raw_end
            chunk = normalized[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start = end
            while start < len(normalized) and normalized[start].isspace():
                start += 1
        return chunks

    def _build_provider_prompt(
        self,
        *,
        official_prompt: str,
        file_content: str,
        execution_profile: ExecutionOperationalProfile,
    ) -> str:
        normalized_content = self._normalize_prompt_block_text(
            file_content,
            dedupe_consecutive_lines=True,
        )
        context_block = f"Arquivo de entrada para analise:\n{normalized_content or '(arquivo sem conteudo textual)'}"
        return self._assemble_prompt(
            instruction_text=str(official_prompt or ""),
            row_data=None,
            auxiliary_context=context_block,
            execution_profile=execution_profile,
        )

    def _enforce_token_limit(
        self,
        *,
        prompt: str,
        provider_runtime: ProviderRuntimeSelection,
    ) -> tuple[str, bool]:
        max_tokens_allowed = max(settings.max_tokens_per_execution, 1)
        current_prompt = prompt
        current_tokens = provider_runtime.client.count_tokens(current_prompt)
        if current_tokens <= max_tokens_allowed:
            return current_prompt, False

        was_truncated = False
        for _ in range(12):
            ratio = max_tokens_allowed / max(current_tokens, 1)
            new_length = max(500, int(len(current_prompt) * ratio * 0.9))
            if new_length >= len(current_prompt):
                new_length = len(current_prompt) - 1
            if new_length <= 0:
                break
            truncated_prompt, _ = self._truncate_text_at_boundary(
                text=current_prompt,
                max_characters=new_length,
                truncation_notice=None,
            )
            if len(truncated_prompt) >= len(current_prompt):
                truncated_prompt = current_prompt[:new_length].rstrip()
            current_prompt = truncated_prompt
            if not current_prompt:
                break
            current_tokens = provider_runtime.client.count_tokens(current_prompt)
            was_truncated = True
            if current_tokens <= max_tokens_allowed:
                return current_prompt, was_truncated

        raise AppException(
            "Prompt exceeds configured token limit for execution.",
            status_code=422,
            code="prompt_token_limit_exceeded",
            details={"max_tokens_per_execution": settings.max_tokens_per_execution},
        )

    def _execute_with_runtime(
        self,
        *,
        prompt_input: str,
        runtime: ProviderRuntimeSelection,
        client_request_id: str | None = None,
    ) -> ProviderExecutionResult:
        raw_model_metadata = getattr(runtime.model, "config_json", None)
        model_metadata: dict[str, Any] = dict(raw_model_metadata) if isinstance(raw_model_metadata, dict) else {}
        for key in ("api_family", "token_limit_param", "request_profile", "supports_reasoning"):
            if key in model_metadata:
                continue
            raw_value = getattr(runtime.model, key, None)
            if raw_value is not None:
                model_metadata[key] = raw_value
        context_tokens = bind_log_context(
            provider=runtime.provider.slug,
            model=runtime.model.model_slug,
        )
        try:
            return runtime.client.execute_prompt(
                prompt=prompt_input,
                model_name=runtime.model.model_slug,
                max_tokens=settings.max_tokens,
                temperature=settings.temperature,
                model_metadata=model_metadata or None,
                client_request_id=client_request_id,
            )
        finally:
            reset_log_context(context_tokens)

    def _is_concurrency_limited(self, *, queue_job_id: UUID) -> bool:
        processing = self.queue_jobs.count_processing_jobs(exclude_queue_job_id=queue_job_id)
        return processing >= settings.max_concurrent_executions

    def _schedule_retry(
        self,
        *,
        execution_id: UUID,
        queue_job: DjangoAiQueueJob,
        reason: str,
        worker_name: str,
        correlation_id: str | None,
        error_diagnostic: ExecutionErrorDiagnostic | None = None,
    ) -> None:
        current_retry = queue_job.retry_count or 0
        hard_retry_limit = self._safe_hard_limit(settings.max_job_retries_hard_limit, fallback=1)
        if current_retry >= hard_retry_limit:
            hard_retry_exception = AppException(
                "Execution exceeded hard retry limit.",
                status_code=422,
                code="job_retries_hard_limit_exceeded",
                details={
                    "execution_id": str(execution_id),
                    "current_retry": int(current_retry),
                    "max_job_retries_hard_limit": hard_retry_limit,
                },
            )
            hard_retry_diagnostic = classify_execution_error(
                hard_retry_exception,
                failure_phase="execution.process.retry_gate",
            )
            self._mark_execution_failed(
                execution_id=execution_id,
                queue_job_id=queue_job.id,
                error_message=hard_retry_exception.payload.message,
                worker_name=worker_name,
                ip_address=None,
                register_error_file=True,
                error_diagnostic=hard_retry_diagnostic,
            )
            return

        if current_retry >= settings.max_retries:
            self._mark_execution_failed(
                execution_id=execution_id,
                queue_job_id=queue_job.id,
                error_message=reason,
                worker_name=worker_name,
                ip_address=None,
                register_error_file=True,
                error_diagnostic=error_diagnostic,
            )
            return

        next_retry = current_retry + 1
        backoff_base = settings.retry_backoff or settings.retry_backoff_seconds
        delay_seconds = backoff_base * (2 ** (next_retry - 1))
        delay_ms = delay_seconds * 1000

        self.operational_session.rollback()
        self.shared_session.rollback()
        self.queue_jobs.mark_queued_for_retry(
            queue_job_id=queue_job.id,
            retry_count=next_retry,
            error_message=reason,
        )
        self.shared_executions.update_status(execution_id=execution_id, status=ExecutionStatus.QUEUED.value)
        self.audit_logs.add(
            DjangoAiAuditLog(
                action_type="execution_retry_scheduled",
                entity_type="analysis_executions",
                entity_id=str(execution_id),
                performed_by_user_id=None,
                changes_json={
                    "queue_job_id": str(queue_job.id),
                    "retry_attempt": next_retry,
                    "max_retries": settings.max_retries,
                    "delay_seconds": delay_seconds,
                    "reason": reason[:200],
                    "failure_phase": error_diagnostic.failure_phase if error_diagnostic else None,
                    "error_code": error_diagnostic.error_code if error_diagnostic else None,
                    "error_category": error_diagnostic.error_category if error_diagnostic else None,
                },
                ip_address=None,
            )
        )
        self.operational_session.commit()
        self.shared_session.commit()
        logger.warning(
            "Retry scheduled for execution.",
            extra={
                "execution_id": str(execution_id),
                "event": "retry_attempt",
                "status": "queued",
                "queue_job_id": str(queue_job.id),
                "retry_attempt": next_retry,
                "delay_seconds": delay_seconds,
                "phase": error_diagnostic.failure_phase if error_diagnostic else "execution.process.retry",
                "error_code": error_diagnostic.error_code if error_diagnostic else None,
                "error_category": error_diagnostic.error_category if error_diagnostic else None,
            },
        )
        enqueue_execution_job(
            execution_id=execution_id,
            queue_job_id=queue_job.id,
            correlation_id=correlation_id,
            delay_ms=delay_ms,
        )

    def _should_retry(self, *, exc: Exception, retry_count: int) -> bool:
        if retry_count >= settings.max_retries:
            return False
        if self._is_retryable_provider_exception(exc):
            return True
        if isinstance(exc, (TimeoutError, ConnectionError)):
            return True
        return False

    def _is_retryable_provider_exception(self, exc: Exception) -> bool:
        if not isinstance(exc, AppException):
            return False
        if exc.payload.code in RETRYABLE_ERROR_CODES:
            return True
        if exc.payload.code == "provider_http_error":
            status_code = (exc.payload.details or {}).get("status_code")
            try:
                status_code_int = int(status_code)
            except (TypeError, ValueError):
                return False
            return status_code_int in RETRYABLE_PROVIDER_STATUS_CODES
        return False

    @staticmethod
    def _is_fatal_tabular_row_exception(exc: Exception) -> bool:
        if not isinstance(exc, AppException):
            return False
        return exc.payload.code in FATAL_TABULAR_ERROR_CODES

    @staticmethod
    def _error_message(exc: Exception) -> str:
        if isinstance(exc, AppException):
            if exc.payload.code in {"provider_http_error", "provider_timeout", "provider_network_error"}:
                details = exc.payload.details if isinstance(exc.payload.details, dict) else {}
                if details:
                    return summarize_provider_error_message(details=details)
            return exc.payload.message
        return str(exc)

    @staticmethod
    def _normalize_prompt_override(value: str | None) -> str | None:
        normalized = str(value or "").strip()
        if not normalized:
            return None
        return normalized

    def _mark_execution_failed(
        self,
        *,
        execution_id: UUID,
        queue_job_id: UUID,
        error_message: str,
        worker_name: str,
        ip_address: str | None,
        register_error_file: bool,
        error_diagnostic: ExecutionErrorDiagnostic | None = None,
    ) -> None:
        self.operational_session.rollback()
        self.shared_session.rollback()

        queue_job = self.queue_jobs.get_by_id(queue_job_id)
        if queue_job is not None:
            queue_job.job_status = ExecutionStatus.FAILED.value
            queue_job.error_message = error_message[:2000]
            queue_job.worker_name = worker_name
            queue_job.finished_at = datetime.now(timezone.utc)
            queue_job.retry_count = max(queue_job.retry_count or 0, 0)

        if register_error_file:
            self._register_error_file(execution_id=execution_id, error_message=error_message)

        self.shared_executions.update_status(execution_id=execution_id, status=ExecutionStatus.FAILED.value)
        self.audit_logs.add(
            DjangoAiAuditLog(
                action_type="execution_failed",
                entity_type="analysis_executions",
                entity_id=str(execution_id),
                performed_by_user_id=None,
                changes_json={
                    "queue_job_id": str(queue_job_id),
                    "error_message": error_message[:500],
                    "failure_phase": error_diagnostic.failure_phase if error_diagnostic else None,
                    "error_code": error_diagnostic.error_code if error_diagnostic else None,
                    "error_category": error_diagnostic.error_category if error_diagnostic else None,
                },
                ip_address=ip_address,
            )
        )

        self.operational_session.commit()
        self.shared_session.commit()
        logger.error(
            "Execution marked as failed.",
            extra={
                "execution_id": str(execution_id),
                "event": "execution_failed",
                "queue_job_id": str(queue_job_id),
                "phase": error_diagnostic.failure_phase if error_diagnostic else "execution.process.failed",
                "error_code": error_diagnostic.error_code if error_diagnostic else None,
                "error_category": error_diagnostic.error_category if error_diagnostic else None,
            },
        )

    def _register_error_file(self, *, execution_id: UUID, error_message: str) -> None:
        try:
            self.file_service.register_generated_execution_file(
                execution_id=execution_id,
                file_type="error",
                file_name=f"execution_{execution_id}_error.txt",
                content=error_message[:5000].encode("utf-8"),
                mime_type="text/plain",
            )
        except Exception:
            logger.exception(
                "Failed to register execution error file.",
                extra={"execution_id": str(execution_id)},
            )

    @staticmethod
    def _parse_status(raw_status: str) -> ExecutionStatus:
        try:
            return ExecutionStatus(raw_status)
        except ValueError:
            return ExecutionStatus.FAILED
