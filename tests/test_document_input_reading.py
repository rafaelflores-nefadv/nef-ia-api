import io

from docx import Document

from app.services.execution_service import ExecutionService


def test_extract_docx_text_reads_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Resumo do contrato")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Cliente"
    table.cell(0, 1).text = "NabeFerro"
    buffer = io.BytesIO()
    document.save(buffer)

    text = ExecutionService._extract_docx_text(buffer.getvalue())

    assert "Resumo do contrato" in text
    assert "Cliente | NabeFerro" in text
